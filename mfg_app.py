#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Manufacturing Orders Management System - Laser/Press
Version: 1.1
Author: Production IT Team
Framework: CustomTkinter + Supabase
"""

import os
import sys
import json
import uuid
import datetime
import threading
import webbrowser
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, asdict
from enum import Enum

# GUI imports
import customtkinter as ctk
from tkinter import messagebox, filedialog
import tkinter as tk
from PIL import Image, ImageTk

# Data processing
import pandas as pd
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

# Charts
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import matplotlib.dates as mdates

# Database
from supabase import create_client, Client
from dotenv import load_dotenv

# Date picker
from tkcalendar import DateEntry

# Import enhanced customer module
from customer_module_enhanced import (
    CustomerExtended,
    CustomerValidator,
    CustomerEditDialog,
    CustomerSearchDialog,
    CustomerExportDialog
)

# Import attachments and WZ modules
from order_attachments_widget import OrderAttachmentsWidget  # Nowy widget z SERVICE_ROLE
from order_confirmation_dialog import OrderConfirmationDialog
from wz_dialog import WZGeneratorDialog

# Import settings modules
from settings_manager import get_settings_manager, initialize_settings
from settings_dialog import SettingsDialog

# Import thumbnail loader
from thumbnail_loader import get_thumbnail_loader

# Import treeview utilities
from treeview_utils import ResizableTreeView, configure_treeview_style, enable_column_sorting

# Load environment variables
load_dotenv()

# Initialize settings manager
initialize_settings()

# Configure CustomTkinter
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

# Constants
APP_NAME = "System Zarządzania Produkcją - Laser/Prasa"
VERSION = "1.1"
WINDOW_WIDTH = 1400
WINDOW_HEIGHT = 800

# Colors for status
STATUS_COLORS = {
    'RECEIVED': '#FFA500',      # Orange
    'CONFIRMED': '#4169E1',     # Royal Blue
    'PLANNED': '#9370DB',       # Medium Purple
    'IN_PROGRESS': '#FFD700',   # Gold
    'DONE': '#32CD32',          # Lime Green
    'INVOICED': '#808080'       # Gray
}

STATUS_NAMES = {
    'RECEIVED': 'Wpłynęło',
    'CONFIRMED': 'Potwierdzono',
    'PLANNED': 'Na planie',
    'IN_PROGRESS': 'W realizacji',
    'DONE': 'Gotowe',
    'INVOICED': 'Wyfakturowane'
}

class OrderStatus(Enum):
    RECEIVED = 'RECEIVED'
    CONFIRMED = 'CONFIRMED'
    PLANNED = 'PLANNED'
    IN_PROGRESS = 'IN_PROGRESS'
    DONE = 'DONE'
    INVOICED = 'INVOICED'

# Using CustomerExtended from customer_module_enhanced instead of basic Customer
# The CustomerExtended class provides all fields including NIP, REGON, address, etc.

@dataclass
class Order:
    id: Optional[str] = None
    process_no: str = ""
    customer_id: str = ""
    title: str = ""
    status: str = "RECEIVED"
    price_pln: float = 0.0
    received_at: Optional[str] = None
    planned_at: Optional[str] = None
    finished_at: Optional[str] = None
    notes: str = ""
    created_at: Optional[str] = None
    customer_name: Optional[str] = None

@dataclass
class Part:
    id: Optional[str] = None
    order_id: str = ""
    idx_code: str = ""
    name: str = ""
    material: str = ""
    thickness_mm: float = 0.0
    qty: int = 1
    unit_price: float = 0.0

class SupabaseManager:
    """Manager for all Supabase operations with user authentication"""

    def __init__(self):
        # Używaj centralnego klienta z SERVICE_ROLE_KEY
        from supabase_client import get_supabase_client
        self.client: Client = get_supabase_client()
        print("[OK] SupabaseManager używa centralnego klienta (SERVICE_ROLE)")
    
    # Customer operations
    def get_customers(self) -> List[Dict]:
        try:
            response = self.client.table('customers').select("*").order('name').execute()
            return response.data
        except Exception as e:
            print(f"Error getting customers: {e}")
            return []
    
    def create_customer(self, customer: CustomerExtended) -> Optional[Dict]:
        try:
            # Map CustomerExtended fields to database columns
            data = customer.to_dict() if hasattr(customer, 'to_dict') else {
                'name': customer.name,
                'short_name': getattr(customer, 'short_name', None),
                'nip': getattr(customer, 'nip', None),
                'regon': getattr(customer, 'regon', None),
                'krs': getattr(customer, 'krs', None),
                'email': getattr(customer, 'email', None),
                'phone': getattr(customer, 'phone', None),
                'website': getattr(customer, 'website', None),
                'address': getattr(customer, 'address', None),
                'city': getattr(customer, 'city', None),
                'postal_code': getattr(customer, 'postal_code', None),
                'country': getattr(customer, 'country', 'Polska'),
                'contact_person': getattr(customer, 'contact_person', None),
                'contact_position': getattr(customer, 'contact_position', None),
                'contact_phone': getattr(customer, 'contact_phone', None),
                'contact_email': getattr(customer, 'contact_email', None),
                'credit_limit': getattr(customer, 'credit_limit', 0),
                'payment_terms': getattr(customer, 'payment_terms', 14),
                'discount_percent': getattr(customer, 'discount_percent', 0),
                'notes': getattr(customer, 'notes', None),
                'tags': getattr(customer, 'tags', []),
                'is_active': getattr(customer, 'is_active', True),
                'customer_type': getattr(customer, 'customer_type', 'company')
            }

            # Clean empty strings and empty lists - convert to None to satisfy database constraints
            # Database constraints require NULL instead of empty strings for email fields
            cleaned_data = {}
            for key, value in data.items():
                if value == '' or value == []:
                    # Convert empty strings and empty lists to None
                    cleaned_data[key] = None
                elif isinstance(value, str):
                    # Strip whitespace from strings
                    cleaned_data[key] = value.strip() if value.strip() else None
                else:
                    cleaned_data[key] = value

            # Clean NIP/REGON - remove dashes and spaces
            if cleaned_data.get('nip'):
                cleaned_data['nip'] = cleaned_data['nip'].replace('-', '').replace(' ', '')
            if cleaned_data.get('regon'):
                cleaned_data['regon'] = cleaned_data['regon'].replace('-', '').replace(' ', '')

            response = self.client.table('customers').insert(cleaned_data).execute()
            return response.data[0] if response.data else None
        except Exception as e:
            error_message = str(e)
            print(f"Error creating customer: {e}")

            # Provide specific error messages for common issues
            if 'check_contact_email_format' in error_message:
                print("❌ Nieprawidłowy format contact_email - sprawdź adres email osoby kontaktowej")
            elif 'check_email_format' in error_message:
                print("❌ Nieprawidłowy format email - sprawdź główny adres email")
            elif 'duplicate key' in error_message and 'nip' in error_message:
                print("❌ Klient z tym numerem NIP już istnieje w bazie")
            elif 'Invalid NIP' in error_message:
                print("❌ Nieprawidłowy numer NIP - sprawdź poprawność numeru")
            elif 'Invalid REGON' in error_message:
                print("❌ Nieprawidłowy numer REGON - sprawdź poprawność numeru")

            return None
    
    def update_customer(self, customer_id: str, updates: Dict) -> bool:
        try:
            # Clean empty strings and empty lists - convert to None to satisfy database constraints
            cleaned_updates = {}
            for key, value in updates.items():
                if value == '' or value == []:
                    # Convert empty strings and empty lists to None
                    cleaned_updates[key] = None
                elif isinstance(value, str):
                    # Strip whitespace from strings
                    cleaned_updates[key] = value.strip() if value.strip() else None
                else:
                    cleaned_updates[key] = value

            # Clean NIP/REGON - remove dashes and spaces
            if cleaned_updates.get('nip'):
                cleaned_updates['nip'] = cleaned_updates['nip'].replace('-', '').replace(' ', '')
            if cleaned_updates.get('regon'):
                cleaned_updates['regon'] = cleaned_updates['regon'].replace('-', '').replace(' ', '')

            self.client.table('customers').update(cleaned_updates).eq('id', customer_id).execute()
            return True
        except Exception as e:
            error_message = str(e)
            print(f"Error updating customer: {e}")

            # Provide specific error messages for common issues
            if 'check_contact_email_format' in error_message:
                print("❌ Nieprawidłowy format contact_email - sprawdź adres email osoby kontaktowej")
            elif 'check_email_format' in error_message:
                print("❌ Nieprawidłowy format email - sprawdź główny adres email")
            elif 'duplicate key' in error_message and 'nip' in error_message:
                print("❌ Klient z tym numerem NIP już istnieje w bazie")

            return False
    
    def delete_customer(self, customer_id: str) -> bool:
        try:
            self.client.table('customers').delete().eq('id', customer_id).execute()
            return True
        except Exception as e:
            print(f"Error deleting customer: {e}")
            return False
    
    # Order operations
    def get_orders(self, filters: Dict = None) -> List[Dict]:
        try:
            query = self.client.table('v_orders_full').select("*")
            
            if filters:
                if filters.get('customer_id'):
                    query = query.eq('customer_id', filters['customer_id'])
                if filters.get('status'):
                    query = query.eq('status', filters['status'])
                if filters.get('date_from'):
                    query = query.gte('received_at', filters['date_from'])
                if filters.get('date_to'):
                    query = query.lte('received_at', filters['date_to'])
                if filters.get('title'):
                    query = query.ilike('title', f"%{filters['title']}%")
            
            response = query.order('created_at', desc=True).execute()
            return response.data
        except Exception as e:
            print(f"Error getting orders: {e}")
            return []
    
    def create_order(self, order: Order) -> Optional[Dict]:
        try:
            data = {
                'customer_id': order.customer_id,
                'title': order.title,
                'status': order.status,
                'price_pln': order.price_pln,
                'received_at': order.received_at,
                'planned_at': order.planned_at,
                'finished_at': order.finished_at,
                'notes': order.notes
            }
            response = self.client.table('orders').insert(data).execute()
            return response.data[0] if response.data else None
        except Exception as e:
            print(f"Error creating order: {e}")
            return None
    
    def update_order(self, order_id: str, updates: Dict) -> bool:
        try:
            self.client.table('orders').update(updates).eq('id', order_id).execute()
            return True
        except Exception as e:
            print(f"Error updating order: {e}")
            return False
    
    def delete_order(self, order_id: str) -> bool:
        try:
            # First delete all order items
            self.client.table('order_items').delete().eq('order_id', order_id).execute()
            # Then delete order
            self.client.table('orders').delete().eq('id', order_id).execute()
            return True
        except Exception as e:
            print(f"Error deleting order: {e}")
            return False
    
    # Parts operations
    def get_parts(self, order_id: str) -> List[Dict]:
        try:
            # Pobierz części z tabeli order_items zamiast parts
            # UWAGA: Zmieniono z 'parts' na 'order_items' zgodnie ze strukturą bazy danych
            response = self.client.table('order_items').select(
                "*",
                "products_catalog(preview_800_url, thumbnail_100_url)"
            ).eq('order_id', order_id).execute()

            # Przetwórz dane aby dodać miniatury do części
            parts = response.data
            for part in parts:
                # Jeśli część ma powiązany produkt, dodaj miniaturę
                if part.get('products_catalog'):
                    product = part['products_catalog']
                    if product.get('preview_800_url'):
                        part['preview_800_url'] = product['preview_800_url']
                    if product.get('thumbnail_100_url'):
                        part['thumbnail_100_url'] = product['thumbnail_100_url']
                    # Usuń zagnieżdżony obiekt produktu po skopiowaniu danych
                    del part['products_catalog']

            return parts
        except Exception as e:
            print(f"Error getting parts: {e}")
            return []
    
    def create_part(self, part: Part) -> Optional[Dict]:
        try:
            # Mapowanie pól z Part na order_items
            data = {
                'order_id': part.order_id,
                'product_id': part.id if part.id else str(uuid.uuid4()),  # Użyj ID produktu lub wygeneruj nowe
                'product_name': part.name,  # Zmiana z 'name' na 'product_name'
                'quantity': part.qty,  # Zmiana z 'qty' na 'quantity'
                'unit_price': part.unit_price,
                'total_price': part.qty * part.unit_price,
                'notes': f"IDX: {part.idx_code}, Material: {part.material}, Grubość: {part.thickness_mm}mm"
            }
            response = self.client.table('order_items').insert(data).execute()
            return response.data[0] if response.data else None
        except Exception as e:
            print(f"Error creating part: {e}")
            return None
    
    def update_part(self, part_id: str, updates: Dict) -> bool:
        try:
            # Mapowanie pól jeśli potrzeba
            if 'qty' in updates:
                updates['quantity'] = updates.pop('qty')
            if 'name' in updates:
                updates['product_name'] = updates.pop('name')

            self.client.table('order_items').update(updates).eq('id', part_id).execute()
            return True
        except Exception as e:
            print(f"Error updating part: {e}")
            return False

    def delete_part(self, part_id: str) -> bool:
        try:
            self.client.table('order_items').delete().eq('id', part_id).execute()
            return True
        except Exception as e:
            print(f"Error deleting part: {e}")
            return False
    
    # File operations
    def upload_file(self, order_id: str, process_no: str, file_path: str) -> bool:
        try:
            filename = os.path.basename(file_path)
            storage_path = f"orders/{process_no}/{filename}"
            
            with open(file_path, 'rb') as f:
                self.client.storage.from_('attachments').upload(
                    path=storage_path,
                    file=f,
                    file_options={"content-type": "application/octet-stream"}
                )
            
            # Record in database
            file_data = {
                'order_id': order_id,
                'storage_path': storage_path,
                'original_name': filename
            }
            self.client.table('files').insert(file_data).execute()
            return True
        except Exception as e:
            print(f"Error uploading file: {e}")
            return False
    
    def get_files(self, order_id: str) -> List[Dict]:
        try:
            response = self.client.table('files').select("*").eq('order_id', order_id).execute()
            return response.data
        except Exception as e:
            print(f"Error getting files: {e}")
            return []
    
    # Dashboard operations
    def get_status_counts(self) -> Dict[str, int]:
        try:
            response = self.client.table('v_orders_status_counts').select("*").execute()
            counts = {item['status']: item['cnt'] for item in response.data}
            return counts
        except Exception as e:
            print(f"Error getting status counts: {e}")
            return {}
    
    def get_sla_dashboard(self) -> Dict[str, int]:
        try:
            response = self.client.table('v_orders_sla').select("*").execute()

            overdue = 0
            soon = 0
            on_time = 0

            for order in response.data:
                if order.get('overdue'):
                    overdue += 1
                elif order.get('days_to_deadline') is not None:
                    if order['days_to_deadline'] <= 2:
                        soon += 1
                    else:
                        on_time += 1

            return {
                'overdue': overdue,
                'soon': soon,
                'on_time': on_time
            }
        except Exception as e:
            print(f"Error getting SLA dashboard: {e}")
            return {'overdue': 0, 'soon': 0, 'on_time': 0}

    # New methods for enhanced customer module
    def search_customers(self, criteria: Dict) -> List[Dict]:
        """Advanced customer search with multiple criteria"""
        try:
            query = self.client.table('customers').select("*")

            if criteria.get('name'):
                query = query.ilike('name', f"%{criteria['name']}%")
            if criteria.get('short_name'):
                query = query.ilike('short_name', f"%{criteria['short_name']}%")
            if criteria.get('nip'):
                query = query.eq('nip', criteria['nip'])
            if criteria.get('regon'):
                query = query.eq('regon', criteria['regon'])
            if criteria.get('city'):
                query = query.ilike('city', f"%{criteria['city']}%")
            if criteria.get('contact_person'):
                query = query.ilike('contact_person', f"%{criteria['contact_person']}%")
            if criteria.get('is_active') is not None:
                query = query.eq('is_active', criteria['is_active'])
            if criteria.get('tags') and len(criteria['tags']) > 0:
                # Search for customers with any of the specified tags
                query = query.contains('tags', criteria['tags'])

            response = query.order('name').execute()
            return response.data
        except Exception as e:
            print(f"Error searching customers: {e}")
            return []

    def get_customer_statistics(self, customer_id: str) -> Dict:
        """Get customer statistics including orders and financial data"""
        try:
            # Use view if available, otherwise calculate manually
            response = self.client.table('v_customer_statistics').select("*").eq('id', customer_id).execute()
            if response.data:
                return response.data[0]

            # Fallback: calculate manually
            orders_response = self.client.table('orders').select("*").eq('customer_id', customer_id).execute()
            orders = orders_response.data

            total_orders = len(orders)
            total_revenue = sum(float(o.get('price_pln', 0) or 0) for o in orders)
            avg_order_value = total_revenue / total_orders if total_orders > 0 else 0

            # Get last order date
            last_order_date = None
            if orders:
                sorted_orders = sorted(orders, key=lambda x: x.get('created_at', ''), reverse=True)
                last_order_date = sorted_orders[0].get('created_at')

            return {
                'order_count': total_orders,
                'total_revenue': total_revenue,
                'avg_order_value': avg_order_value,
                'last_order_date': last_order_date
            }
        except Exception as e:
            print(f"Error getting customer statistics: {e}")
            return {}

# CustomerDialog - Main customer management window using enhanced modules
class CustomerDialog(ctk.CTkToplevel):
    """Main dialog for managing customers with enhanced features"""

    def __init__(self, parent, db: SupabaseManager):
        super().__init__(parent)
        self.db = db
        self.parent = parent

        self.title("Zarządzanie Klientami - Rozszerzone")
        self.geometry("1200x700")

        # Make modal
        self.transient(parent)
        self.grab_set()

        self.setup_ui()
        self.load_customers()

        # Center window
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (1200 // 2)
        y = (self.winfo_screenheight() // 2) - (700 // 2)
        self.geometry(f"+{x}+{y}")

    def setup_ui(self):
        # Main container
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Header
        header_frame = ctk.CTkFrame(main_frame)
        header_frame.pack(fill="x", padx=5, pady=5)

        ctk.CTkLabel(
            header_frame,
            text="Baza Klientów - Pełne Dane",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(side="left", padx=10)

        # Buttons
        btn_frame = ctk.CTkFrame(header_frame)
        btn_frame.pack(side="right", padx=10)

        ctk.CTkButton(
            btn_frame,
            text="➕ Dodaj",
            width=100,
            command=self.add_customer
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="✏️ Edytuj",
            width=100,
            command=self.edit_customer
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="🔍 Szukaj",
            width=100,
            command=self.search_customers
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="📥 Eksport",
            width=100,
            command=self.export_customers
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="🗑️ Usuń",
            width=100,
            command=self.delete_customer
        ).pack(side="left", padx=5)

        # Customer list
        list_frame = ctk.CTkFrame(main_frame)
        list_frame.pack(fill="both", expand=True, padx=5, pady=5)

        # Create Treeview for customers with enhanced columns
        from tkinter import ttk

        self.tree = ttk.Treeview(
            list_frame,
            columns=('name', 'nip', 'email', 'phone', 'city', 'contact_person', 'credit_limit', 'active'),
            show='headings',
            selectmode='browse'
        )

        self.tree.heading('name', text='Nazwa firmy')
        self.tree.heading('nip', text='NIP')
        self.tree.heading('email', text='Email')
        self.tree.heading('phone', text='Telefon')
        self.tree.heading('city', text='Miasto')
        self.tree.heading('contact_person', text='Osoba kontaktowa')
        self.tree.heading('credit_limit', text='Limit kredytowy')
        self.tree.heading('active', text='Aktywny')

        self.tree.column('name', width=200)
        self.tree.column('nip', width=120)
        self.tree.column('email', width=180)
        self.tree.column('phone', width=120)
        self.tree.column('city', width=100)
        self.tree.column('contact_person', width=150)
        self.tree.column('credit_limit', width=100)
        self.tree.column('active', width=60)

        # Scrollbars
        v_scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.tree.yview)
        h_scrollbar = ttk.Scrollbar(list_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

        self.tree.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")

        list_frame.grid_columnconfigure(0, weight=1)
        list_frame.grid_rowconfigure(0, weight=1)

        # Style for treeview
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Treeview", background="#212121", foreground="white",
                       fieldbackground="#212121", borderwidth=0)
        style.configure("Treeview.Heading", background="#313131", foreground="white")
        style.map('Treeview', background=[('selected', '#144870')])

        # Bottom info panel
        info_frame = ctk.CTkFrame(main_frame)
        info_frame.pack(fill="x", padx=5, pady=5)

        self.info_label = ctk.CTkLabel(
            info_frame,
            text="Wybierz klienta, aby zobaczyć szczegóły",
            font=ctk.CTkFont(size=12)
        )
        self.info_label.pack(side="left", padx=10)

        # Bind double-click to edit
        self.tree.bind("<Double-1>", lambda e: self.edit_customer())

    def load_customers(self):
        """Load customers from database with enhanced data"""
        self.tree.delete(*self.tree.get_children())

        customers = self.db.get_customers()
        for customer in customers:
            # Format credit limit
            credit_limit = f"{customer.get('credit_limit', 0):.0f} PLN" if customer.get('credit_limit') else "0 PLN"

            # Format active status
            is_active = "Tak" if customer.get('is_active', True) else "Nie"

            self.tree.insert('', 'end',
                           values=(
                               customer.get('name', ''),
                               customer.get('nip', ''),
                               customer.get('email', ''),
                               customer.get('phone', ''),
                               customer.get('city', ''),
                               customer.get('contact_person', ''),
                               credit_limit,
                               is_active
                           ),
                           tags=(customer['id'],))

        # Update info
        count = len(customers)
        self.info_label.configure(text=f"Liczba klientów: {count}")

    def add_customer(self):
        """Add new customer using enhanced dialog"""
        dialog = CustomerEditDialog(self, self.db)
        self.wait_window(dialog)
        self.load_customers()

    def edit_customer(self):
        """Edit selected customer using enhanced dialog"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz klienta do edycji")
            return

        item = self.tree.item(selection[0])
        customer_id = item['tags'][0]

        # Get full customer data
        customers = self.db.get_customers()
        customer_data = None
        for customer in customers:
            if customer.get('id') == customer_id:
                customer_data = customer
                break

        if customer_data:
            dialog = CustomerEditDialog(self, self.db, customer_data)
            self.wait_window(dialog)
            self.load_customers()

    def search_customers(self):
        """Open search dialog"""
        dialog = CustomerSearchDialog(self, self.db)
        self.wait_window(dialog)
        # Refresh list after search (search dialog may filter results)
        if hasattr(dialog, 'selected_customer'):
            # Optionally highlight selected customer
            pass

    def export_customers(self):
        """Open export dialog"""
        customers = self.db.get_customers()
        if not customers:
            messagebox.showwarning("Uwaga", "Brak klientów do eksportu")
            return
        dialog = CustomerExportDialog(self, customers)
        self.wait_window(dialog)

    def delete_customer(self):
        """Delete selected customer"""
        selection = self.tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz klienta do usunięcia")
            return

        item = self.tree.item(selection[0])
        customer_id = item['tags'][0]
        customer_name = item['values'][0]

        if messagebox.askyesno("Potwierdzenie",
                               f"Czy na pewno usunąć klienta:\n{customer_name}?"):
            if self.db.delete_customer(customer_id):
                messagebox.showinfo("Sukces", "Klient został usunięty")
                self.load_customers()
            else:
                messagebox.showerror("Błąd", "Nie udało się usunąć klienta")

class OrderDialog(ctk.CTkToplevel):
    """Dialog for creating/editing orders"""
    
    def __init__(self, parent, db: SupabaseManager, order_data=None):
        super().__init__(parent)
        self.db = db
        self.order_data = order_data
        self.parts_list = []
        self.order_id = order_data['id'] if order_data else None

        self.title("Edycja zamówienia" if order_data else "Nowe zamówienie")
        self.geometry("1000x700")
        
        # Make modal
        self.transient(parent)
        self.grab_set()
        
        self.setup_ui()
        
        if order_data:
            self.load_order_data()
        
        # Center window
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (1000 // 2)
        y = (self.winfo_screenheight() // 2) - (700 // 2)
        self.geometry(f"+{x}+{y}")
    
    def setup_ui(self):
        # Main container with scrollable frame
        main_frame = ctk.CTkScrollableFrame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Header
        header_frame = ctk.CTkFrame(main_frame)
        header_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(
            header_frame,
            text="Edycja zamówienia" if self.order_data else "Nowe zamówienie",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(side="left", padx=10)
        
        # Order details section
        order_frame = ctk.CTkFrame(main_frame)
        order_frame.pack(fill="x", padx=5, pady=10)
        
        ctk.CTkLabel(
            order_frame,
            text="Dane zamówienia",
            font=ctk.CTkFont(size=16, weight="bold")
        ).grid(row=0, column=0, columnspan=4, pady=10, sticky="w")
        
        # Customer selection
        ctk.CTkLabel(order_frame, text="Klient:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        
        customers = self.db.get_customers()
        customer_names = [c['name'] for c in customers]
        self.customer_map = {c['name']: c['id'] for c in customers}
        
        self.customer_combo = ctk.CTkComboBox(
            order_frame,
            values=customer_names,
            width=300
        )
        self.customer_combo.grid(row=1, column=1, padx=5, pady=5)
        
        # Title
        ctk.CTkLabel(order_frame, text="Tytuł zamówienia:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.title_entry = ctk.CTkEntry(order_frame, width=300)
        self.title_entry.grid(row=2, column=1, padx=5, pady=5)
        
        # Price
        ctk.CTkLabel(order_frame, text="Cena (PLN):").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.price_entry = ctk.CTkEntry(order_frame, width=150)
        self.price_entry.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        
        # Dates frame
        dates_frame = ctk.CTkFrame(order_frame)
        dates_frame.grid(row=4, column=0, columnspan=4, pady=10)
        
        # Received date
        ctk.CTkLabel(dates_frame, text="Data wpływu:").pack(side="left", padx=5)
        self.received_date = DateEntry(dates_frame, width=12, background='darkblue',
                                       foreground='white', borderwidth=2)
        self.received_date.pack(side="left", padx=5)
        
        # Planned date
        ctk.CTkLabel(dates_frame, text="Planowana data:").pack(side="left", padx=15)
        self.planned_date = DateEntry(dates_frame, width=12, background='darkblue',
                                     foreground='white', borderwidth=2)
        self.planned_date.pack(side="left", padx=5)
        
        # Finished date
        ctk.CTkLabel(dates_frame, text="Data zakończenia:").pack(side="left", padx=15)
        self.finished_date = DateEntry(dates_frame, width=12, background='darkblue',
                                      foreground='white', borderwidth=2)
        self.finished_date.pack(side="left", padx=5)
        
        # Status
        ctk.CTkLabel(order_frame, text="Status:").grid(row=5, column=0, padx=5, pady=5, sticky="e")
        self.status_combo = ctk.CTkComboBox(
            order_frame,
            values=list(STATUS_NAMES.values()),
            width=200
        )
        self.status_combo.grid(row=5, column=1, padx=5, pady=5, sticky="w")
        self.status_combo.set("Wpłynęło")
        
        # Notes
        ctk.CTkLabel(order_frame, text="Uwagi:").grid(row=6, column=0, padx=5, pady=5, sticky="ne")
        self.notes_text = ctk.CTkTextbox(order_frame, width=400, height=100)
        self.notes_text.grid(row=6, column=1, columnspan=2, padx=5, pady=5)
        
        # Parts section
        parts_frame = ctk.CTkFrame(main_frame)
        parts_frame.pack(fill="both", expand=True, padx=5, pady=10)
        
        parts_header = ctk.CTkFrame(parts_frame)
        parts_header.pack(fill="x", pady=5)
        
        ctk.CTkLabel(
            parts_header,
            text="Detale zamówienia",
            font=ctk.CTkFont(size=16, weight="bold")
        ).pack(side="left", padx=10)
        
        # Parts buttons
        parts_btn_frame = ctk.CTkFrame(parts_header)
        parts_btn_frame.pack(side="right", padx=10)
        
        ctk.CTkButton(
            parts_btn_frame,
            text="📦 Wybierz produkty",
            width=150,
            command=self.add_part,
            fg_color="#9C27B0"
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            parts_btn_frame,
            text="✏️ Edytuj",
            width=100,
            command=self.edit_part
        ).pack(side="left", padx=5)
        
        ctk.CTkButton(
            parts_btn_frame,
            text="🗑️ Usuń",
            width=100,
            command=self.delete_part
        ).pack(side="left", padx=5)
        
        # Parts list with thumbnails
        from tkinter import ttk

        # Info label about inline editing
        edit_info = ctk.CTkLabel(
            parts_frame,
            text="💡 Podwójne kliknięcie na ilość lub cenę umożliwia bezpośrednią edycję",
            font=ctk.CTkFont(size=11),
            text_color="gray"
        )
        edit_info.pack(anchor="w", pady=(5, 5), padx=10)

        # Add checkbox for thumbnails
        show_thumbnails_frame = ctk.CTkFrame(parts_frame)
        show_thumbnails_frame.pack(fill="x", pady=(0, 5))

        self.show_parts_thumbnails_var = tk.BooleanVar(value=True)
        self.show_parts_thumbnails_check = ctk.CTkCheckBox(
            show_thumbnails_frame,
            text="Wyświetlaj miniatury",
            variable=self.show_parts_thumbnails_var,
            command=self.refresh_parts_display
        )
        self.show_parts_thumbnails_check.pack(side="left", padx=10)

        # Use EditablePartsTreeView for inline editing of parts list
        from editable_parts_tree import EditablePartsTreeView

        self.parts_tree = EditablePartsTreeView(
            parts_frame,
            columns=('idx', 'name', 'material', 'thickness', 'qty', 'price'),
            editable_columns=['qty', 'price'],  # Allow editing quantity and price
            on_edit_complete=self._on_part_edit_complete,
            show='tree headings',  # Changed to show tree column for thumbnails
            height=8
        )

        self.parts_tree.heading('#0', text='')  # Thumbnail column
        self.parts_tree.heading('idx', text='Indeks')
        self.parts_tree.heading('name', text='Nazwa')
        self.parts_tree.heading('material', text='Materiał')
        self.parts_tree.heading('thickness', text='Grubość [mm]')
        self.parts_tree.heading('qty', text='Ilość')
        self.parts_tree.heading('price', text='Cena')

        # Set default widths (will be overridden by saved settings if available)
        settings = get_settings_manager().settings
        thumbnail_width = settings.list_row_height if settings.list_show_thumbnails else 50

        self.parts_tree.column('#0', width=thumbnail_width, stretch=False)  # Thumbnail column
        self.parts_tree.column('idx', width=100)
        self.parts_tree.column('name', width=200)
        self.parts_tree.column('material', width=120)
        self.parts_tree.column('thickness', width=80)
        self.parts_tree.column('qty', width=60)
        self.parts_tree.column('price', width=80)

        # Apply global style settings
        configure_treeview_style(self.parts_tree, settings)

        # Enable column sorting
        enable_column_sorting(self.parts_tree,
                            ['idx', 'name', 'material', 'thickness', 'qty', 'price'],
                            numeric_columns=['thickness', 'qty', 'price'])

        self.parts_tree.pack(fill="both", expand=True, pady=5)

        # Store thumbnails to prevent garbage collection
        self.parts_thumbnails = []

        # Załączniki section
        attachments_label = ctk.CTkLabel(
            main_frame,
            text="📎 Załączniki",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        attachments_label.pack(anchor="w", pady=(10, 5))

        # Info for new orders
        if not self.order_id:
            info_label = ctk.CTkLabel(
                main_frame,
                text="ℹ️ Zapisz zamówienie aby móc dodawać załączniki",
                font=ctk.CTkFont(size=11),
                text_color="gray"
            )
            info_label.pack(anchor="w", pady=(0, 5))

        # Załączniki widget - nowy z SERVICE_ROLE_KEY
        self.attachments_widget = OrderAttachmentsWidget(
            main_frame,
            entity_id=self.order_id
        )
        self.attachments_widget.pack(fill="both", expand=True, padx=5, pady=10)

        # Bottom buttons
        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(fill="x", pady=10)

        ctk.CTkButton(
            btn_frame,
            text="💾 Zapisz zamówienie",
            width=200,
            height=40,
            command=self.save_order,
            font=ctk.CTkFont(size=14, weight="bold"),
            fg_color="green"
        ).pack(side="left", padx=10)

        ctk.CTkButton(
            btn_frame,
            text="🖨️ Wydrukuj potwierdzenie",
            width=180,
            height=40,
            command=self.print_confirmation,
            fg_color="blue"
        ).pack(side="left", padx=10)

        ctk.CTkButton(
            btn_frame,
            text="✉️ Wyślij potwierdzenie",
            width=180,
            height=40,
            command=self.send_confirmation,
            fg_color="purple"
        ).pack(side="left", padx=10)

        ctk.CTkButton(
            btn_frame,
            text="Anuluj",
            width=150,
            height=40,
            command=self.destroy,
            fg_color="gray"
        ).pack(side="right", padx=10)
    
    def load_order_data(self):
        """Load existing order data"""
        if not self.order_data:
            return

        # Set customer
        for customer_name, customer_id in self.customer_map.items():
            if customer_id == self.order_data['customer_id']:
                self.customer_combo.set(customer_name)
                break

        # Set fields
        self.title_entry.insert(0, self.order_data.get('title', ''))
        self.price_entry.insert(0, str(self.order_data.get('price_pln', 0)))

        # Set dates
        if self.order_data.get('received_at'):
            self.received_date.set_date(datetime.datetime.fromisoformat(self.order_data['received_at']).date())
        if self.order_data.get('planned_at'):
            self.planned_date.set_date(datetime.datetime.fromisoformat(self.order_data['planned_at']).date())
        if self.order_data.get('finished_at'):
            self.finished_date.set_date(datetime.datetime.fromisoformat(self.order_data['finished_at']).date())

        # Set status
        status_pl = STATUS_NAMES.get(self.order_data.get('status', 'RECEIVED'))
        self.status_combo.set(status_pl)

        # Set notes
        if self.order_data.get('notes'):
            self.notes_text.insert("1.0", self.order_data['notes'])

        # Load parts
        order_id = self.order_data.get('id')
        if order_id:
            parts = self.db.get_parts(order_id)
            self.parts_list = parts
            self.refresh_parts_display()

            # Load attachments for existing order
            self.attachments_widget.load_attachments()
    
    def add_part(self):
        """Add new part - opens product selector dialog"""
        # Import the new product selector
        from products_selector_dialog_v2 import ProductSelectorDialog

        # Open product selector with existing parts
        dialog = ProductSelectorDialog(
            self,
            self.db,
            existing_parts=self.parts_list,
            callback=self.process_selected_products
        )
        self.wait_window(dialog)

    def process_selected_products(self, selected_parts: List[Dict]):
        """Process products selected from selector dialog"""
        # Clear existing parts
        self.parts_list = selected_parts

        # Refresh display with new parts
        self.refresh_parts_display()

        # Calculate and update total price
        total_price = self._calculate_total_price()
        self.price_entry.delete(0, 'end')
        self.price_entry.insert(0, f"{total_price:.2f}")

    def refresh_parts_display(self):
        """Refresh parts tree display with optional thumbnails"""
        # Clear tree
        for item in self.parts_tree.get_children():
            self.parts_tree.delete(item)

        # Clear thumbnails
        self.parts_thumbnails = []

        # Apply current style settings (in case they changed)
        settings = get_settings_manager().settings
        configure_treeview_style(self.parts_tree, settings)

        # Update thumbnail column width based on settings
        thumbnail_width = settings.list_row_height if settings.list_show_thumbnails else 50
        self.parts_tree.column('#0', width=thumbnail_width)

        # Add parts to tree
        for part in self.parts_list:
            # Get thumbnail if enabled
            thumbnail = None
            if self.show_parts_thumbnails_var.get():
                thumbnail = self._get_part_thumbnail(part)
                if thumbnail:
                    self.parts_thumbnails.append(thumbnail)  # Keep reference

            # Get quantity and unit price
            qty = part.get('quantity', part.get('qty', 1))
            unit_price = part.get('unit_price', 0)

            # If no unit price, calculate from costs
            if unit_price == 0:
                unit_price = (
                    part.get('material_laser_cost', 0) +
                    part.get('bending_cost', 0) +
                    part.get('additional_costs', 0)
                )

            # Insert item with thumbnail
            self.parts_tree.insert('', 'end',
                image=thumbnail if thumbnail else '',
                values=(
                    part.get('idx_code', ''),
                    part.get('name', ''),
                    part.get('material', ''),
                    part.get('thickness_mm', ''),
                    qty,
                    f"{unit_price:.2f}"
                )
            )

    def _on_part_edit_complete(self, item_id, column, new_value, old_value):
        """Handle inline edit completion for parts"""
        # Get the item index in the parts_list
        items = self.parts_tree.get_children()
        item_index = items.index(item_id)

        if item_index >= 0 and item_index < len(self.parts_list):
            part = self.parts_list[item_index]

            # Update the part data
            if column == 'qty':
                part['qty'] = int(new_value) if new_value else 0
                part['quantity'] = part['qty']  # Keep both keys in sync
            elif column == 'price':
                part['unit_price'] = float(new_value) if new_value else 0.0

            # Recalculate and update total price
            total_price = self._calculate_total_price()
            self.price_entry.delete(0, 'end')
            self.price_entry.insert(0, f"{total_price:.2f}")

    def _calculate_total_price(self):
        """Calculate total price from parts list"""
        total = 0
        for part in self.parts_list:
            qty = part.get('quantity', part.get('qty', 1))
            unit_price = part.get('unit_price', 0)

            if unit_price == 0:
                unit_price = (
                    part.get('material_laser_cost', 0) +
                    part.get('bending_cost', 0) +
                    part.get('additional_costs', 0)
                )

            total += qty * unit_price
        return total

    def _get_part_thumbnail(self, part):
        """Get thumbnail for part"""
        try:
            # Use global thumbnail loader with row height from settings
            loader = get_thumbnail_loader()
            settings = get_settings_manager().settings
            row_height = settings.list_row_height
            return loader.get_product_thumbnail(part, row_height=row_height)
        except Exception as e:
            print(f"Error loading part thumbnail: {e}")
            return None

    def edit_part(self):
        """Edit selected part"""
        selection = self.parts_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz detal do edycji")
            return
        
        index = self.parts_tree.index(selection[0])
        part_data = self.parts_list[index]
        
        dialog = PartEditDialog(self, self.parts_list, part_data, index)
        self.wait_window(dialog)
        
        if hasattr(dialog, 'part_data'):
            self.parts_list[index] = dialog.part_data
            
            # Update tree
            self.parts_tree.delete(selection[0])
            self.parts_tree.insert('', index, values=(
                dialog.part_data.get('idx_code', ''),
                dialog.part_data['name'],
                dialog.part_data.get('material', ''),
                dialog.part_data.get('thickness_mm', ''),
                dialog.part_data.get('qty', 1)
            ))
    
    def delete_part(self):
        """Delete selected part"""
        selection = self.parts_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz detal do usunięcia")
            return
        
        index = self.parts_tree.index(selection[0])
        part_name = self.parts_list[index]['name']
        
        if messagebox.askyesno("Potwierdzenie", f"Czy usunąć detal '{part_name}'?"):
            del self.parts_list[index]
            self.parts_tree.delete(selection[0])
    
    def save_order(self):
        """Save order to database"""
        # Validate
        customer = self.customer_combo.get()
        if not customer or customer not in self.customer_map:
            messagebox.showwarning("Uwaga", "Wybierz klienta")
            return
        
        title = self.title_entry.get().strip()
        if not title:
            messagebox.showwarning("Uwaga", "Podaj tytuł zamówienia")
            return
        
        try:
            price = float(self.price_entry.get() or 0)
        except ValueError:
            messagebox.showwarning("Uwaga", "Nieprawidłowa cena")
            return
        
        # Get status
        status_pl = self.status_combo.get()
        status = None
        for key, value in STATUS_NAMES.items():
            if value == status_pl:
                status = key
                break
        
        # Prepare order data
        order_data = {
            'customer_id': self.customer_map[customer],
            'title': title,
            'price_pln': price,
            'status': status,
            'received_at': self.received_date.get_date().isoformat(),
            'planned_at': self.planned_date.get_date().isoformat(),
            'finished_at': self.finished_date.get_date().isoformat() if self.finished_date.get_date() else None,
            'notes': self.notes_text.get("1.0", "end-1c").strip()
        }
        
        if self.order_data:
            # Update existing order
            if self.db.update_order(self.order_data['id'], order_data):
                # Update parts (simplified - delete all and recreate)
                existing_parts = self.db.get_parts(self.order_data['id'])
                for part in existing_parts:
                    self.db.delete_part(part['id'])

                # Add new parts
                for part_data in self.parts_list:
                    # Get quantity
                    qty = int(part_data.get('quantity', part_data.get('qty', 1)) or 1)

                    # Get unit price
                    unit_price = float(part_data.get('unit_price', 0) or 0)
                    if unit_price == 0:
                        # Calculate from costs if no direct price
                        unit_price = (
                            float(part_data.get('material_laser_cost', 0) or 0) +
                            float(part_data.get('bending_cost', 0) or 0) +
                            float(part_data.get('additional_costs', 0) or 0)
                        )

                    part = Part(
                        order_id=self.order_data['id'],
                        idx_code=part_data.get('idx_code', ''),
                        name=part_data['name'],
                        material=part_data.get('material', ''),
                        thickness_mm=float(part_data.get('thickness_mm', 0) or 0),
                        qty=qty,
                        unit_price=unit_price
                    )
                    self.db.create_part(part)

                messagebox.showinfo("Sukces", "Zamówienie zostało zaktualizowane")
                # Nie zamykaj okna, aby użytkownik mógł edytować załączniki
                # self.destroy()
            else:
                messagebox.showerror("Błąd", "Nie udało się zaktualizować zamówienia")
        else:
            # Create new order
            order = Order(**order_data)
            result = self.db.create_order(order)
            
            if result:
                # Add parts
                for part_data in self.parts_list:
                    # Get quantity
                    qty = int(part_data.get('quantity', part_data.get('qty', 1)) or 1)

                    # Get unit price
                    unit_price = float(part_data.get('unit_price', 0) or 0)
                    if unit_price == 0:
                        # Calculate from costs if no direct price
                        unit_price = (
                            float(part_data.get('material_laser_cost', 0) or 0) +
                            float(part_data.get('bending_cost', 0) or 0) +
                            float(part_data.get('additional_costs', 0) or 0)
                        )

                    part = Part(
                        order_id=result['id'],
                        idx_code=part_data.get('idx_code', ''),
                        name=part_data['name'],
                        material=part_data.get('material', ''),
                        thickness_mm=float(part_data.get('thickness_mm', 0) or 0),
                        qty=qty,
                        unit_price=unit_price
                    )
                    self.db.create_part(part)

                # Ustaw ID dla widgetu załączników
                self.order_id = result['id']
                self.attachments_widget.set_entity_id(result['id'])

                messagebox.showinfo("Sukces", f"Zamówienie {result['process_no']} zostało utworzone")
                # Nie zamykaj okna, aby użytkownik mógł dodać załączniki
                # self.destroy()
            else:
                messagebox.showerror("Błąd", "Nie udało się utworzyć zamówienia")

    def print_confirmation(self):
        """Print order confirmation"""
        if not self.order_id and not self.order_data:
            messagebox.showwarning("Uwaga", "Najpierw zapisz zamówienie")
            return

        # Open confirmation preview dialog
        dialog = OrderConfirmationDialog(
            self,
            self.db,
            self.order_data or {'id': self.order_id},
            self.parts_list,
            mode='print'
        )
        self.wait_window(dialog)

    def send_confirmation(self):
        """Send order confirmation via email"""
        if not self.order_id and not self.order_data:
            messagebox.showwarning("Uwaga", "Najpierw zapisz zamówienie")
            return

        # Open email sending dialog
        dialog = OrderConfirmationDialog(
            self,
            self.db,
            self.order_data or {'id': self.order_id},
            self.parts_list,
            mode='email'
        )
        self.wait_window(dialog)

class PartEditDialog(ctk.CTkToplevel):
    """Dialog for adding/editing part"""
    
    def __init__(self, parent, parts_list, part_data=None, part_index=None):
        super().__init__(parent)
        self.parts_list = parts_list
        self.part_data_original = part_data
        self.part_index = part_index
        
        self.title("Edycja detalu" if part_data else "Nowy detal")
        self.geometry("500x400")
        
        # Make modal
        self.transient(parent)
        self.grab_set()
        
        self.setup_ui()
        
        # Center window
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (500 // 2)
        y = (self.winfo_screenheight() // 2) - (400 // 2)
        self.geometry(f"+{x}+{y}")
    
    def setup_ui(self):
        # Main container
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        # Index field
        ctk.CTkLabel(main_frame, text="Indeks:", font=ctk.CTkFont(size=14)).pack(anchor="w", pady=5)
        self.idx_entry = ctk.CTkEntry(main_frame, width=400, height=35)
        self.idx_entry.pack(pady=5)
        
        # Name field
        ctk.CTkLabel(main_frame, text="Nazwa detalu*:", font=ctk.CTkFont(size=14)).pack(anchor="w", pady=5)
        self.name_entry = ctk.CTkEntry(main_frame, width=400, height=35)
        self.name_entry.pack(pady=5)
        
        # Material field
        ctk.CTkLabel(main_frame, text="Materiał:", font=ctk.CTkFont(size=14)).pack(anchor="w", pady=5)
        self.material_entry = ctk.CTkEntry(main_frame, width=400, height=35)
        self.material_entry.pack(pady=5)
        
        # Thickness field
        ctk.CTkLabel(main_frame, text="Grubość [mm]:", font=ctk.CTkFont(size=14)).pack(anchor="w", pady=5)
        self.thickness_entry = ctk.CTkEntry(main_frame, width=200, height=35)
        self.thickness_entry.pack(pady=5, anchor="w")
        
        # Quantity field
        ctk.CTkLabel(main_frame, text="Ilość:", font=ctk.CTkFont(size=14)).pack(anchor="w", pady=5)
        self.qty_entry = ctk.CTkEntry(main_frame, width=200, height=35)
        self.qty_entry.pack(pady=5, anchor="w")
        self.qty_entry.insert(0, "1")
        
        # Load data if editing
        if self.part_data_original:
            self.idx_entry.insert(0, self.part_data_original.get('idx_code', ''))
            self.name_entry.insert(0, self.part_data_original['name'])
            self.material_entry.insert(0, self.part_data_original.get('material', ''))
            self.thickness_entry.insert(0, str(self.part_data_original.get('thickness_mm', '')))
            self.qty_entry.insert(0, str(self.part_data_original.get('qty', 1)))
        
        # Buttons
        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(fill="x", pady=20)
        
        ctk.CTkButton(
            btn_frame,
            text="Zapisz",
            width=150,
            command=self.save_part
        ).pack(side="left", padx=10)
        
        ctk.CTkButton(
            btn_frame,
            text="Anuluj",
            width=150,
            command=self.destroy
        ).pack(side="right", padx=10)
    
    def save_part(self):
        """Save part data"""
        name = self.name_entry.get().strip()
        
        if not name:
            messagebox.showwarning("Uwaga", "Nazwa detalu jest wymagana")
            return
        
        # Check for duplicate names (exclude current part if editing)
        for i, part in enumerate(self.parts_list):
            if part['name'] == name and i != self.part_index:
                messagebox.showwarning("Uwaga", f"Detal o nazwie '{name}' już istnieje!")
                return
        
        self.part_data = {
            'idx_code': self.idx_entry.get().strip(),
            'name': name,
            'material': self.material_entry.get().strip(),
            'thickness_mm': self.thickness_entry.get().strip(),
            'qty': self.qty_entry.get().strip() or "1"
        }
        
        self.destroy()

class MainApplication(ctk.CTk):
    """Main application window"""
    
    def __init__(self):
        super().__init__()

        self.title(APP_NAME)
        self.geometry(f"{WINDOW_WIDTH}x{WINDOW_HEIGHT}")

        # Initialize settings manager
        self.settings_manager = get_settings_manager()
        self.settings = self.settings_manager.settings

        # Apply theme settings
        if self.settings.theme_mode:
            ctk.set_appearance_mode(self.settings.theme_mode)
        if self.settings.color_theme:
            ctk.set_default_color_theme(self.settings.color_theme)

        # Initialize database
        try:
            self.db = SupabaseManager()
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie można połączyć z bazą danych:\n{e}")
            sys.exit(1)

        self.setup_ui()
        self.load_orders()
        self.update_dashboard()

        # Center window
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (WINDOW_WIDTH // 2)
        y = (self.winfo_screenheight() // 2) - (WINDOW_HEIGHT // 2)
        self.geometry(f"+{x}+{y}")
        
        # Start auto-refresh
        self.auto_refresh()
    
    def setup_ui(self):
        """Setup main UI"""
        # Configure grid
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # Header
        self.create_header()
        
        # Main content area with sidebar and orders list
        content_frame = ctk.CTkFrame(self)
        content_frame.grid(row=1, column=0, sticky="nsew", padx=5, pady=5)
        content_frame.grid_columnconfigure(1, weight=1)
        content_frame.grid_rowconfigure(0, weight=1)
        
        # Sidebar
        self.create_sidebar(content_frame)
        
        # Orders area
        self.create_orders_area(content_frame)
        
        # Status bar
        self.create_status_bar()
    
    def create_header(self):
        """Create header with title and main buttons"""
        header = ctk.CTkFrame(self)
        header.grid(row=0, column=0, sticky="ew", padx=5, pady=5)

        # Logo - dodaj logo producenta po lewej stronie
        logo_frame = ctk.CTkFrame(header)
        logo_frame.pack(side="left", padx=10)

        # Wczytaj logo
        logo_path = self.settings_manager.get_active_logo_path()
        if logo_path and os.path.exists(logo_path):
            try:
                # Załaduj i przeskaluj logo
                logo_img = Image.open(logo_path)
                logo_img.thumbnail((80, 60), Image.Resampling.LANCZOS)
                self.logo_photo = ImageTk.PhotoImage(logo_img)

                # Wyświetl logo
                logo_label = ctk.CTkLabel(logo_frame, image=self.logo_photo, text="")
                logo_label.pack()
            except Exception as e:
                print(f"Nie można załadować logo: {e}")

        # Title
        title_label = ctk.CTkLabel(
            header,
            text=APP_NAME,
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(side="left", padx=20)

        # Main buttons
        btn_frame = ctk.CTkFrame(header)
        btn_frame.pack(side="right", padx=10)

        ctk.CTkButton(
            btn_frame,
            text="➕ Nowe zamówienie",
            width=150,
            height=35,
            font=ctk.CTkFont(size=14, weight="bold"),
            command=self.new_order
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="📦 Generuj WZ",
            width=130,
            height=35,
            command=self.generate_wz,
            fg_color="#FF6B6B"
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="👥 Klienci",
            width=120,
            height=35,
            command=self.manage_customers
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="📄 Raporty",
            width=120,
            height=35,
            command=self.show_reports
        ).pack(side="left", padx=5)

        ctk.CTkButton(
            btn_frame,
            text="🔄 Odśwież",
            width=100,
            height=35,
            command=self.refresh_all
        ).pack(side="left", padx=5)

        # Przycisk ustawień
        ctk.CTkButton(
            btn_frame,
            text="⚙️ Ustawienia",
            width=110,
            height=35,
            command=self.open_settings,
            fg_color="gray"
        ).pack(side="left", padx=5)
    
    def create_sidebar(self, parent):
        """Create sidebar with filters and dashboard"""
        sidebar = ctk.CTkFrame(parent, width=300)
        sidebar.grid(row=0, column=0, sticky="ns", padx=(0, 5))
        
        # Filters section
        filters_label = ctk.CTkLabel(
            sidebar,
            text="Filtry",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        filters_label.pack(pady=10)
        
        # Customer filter
        ctk.CTkLabel(sidebar, text="Klient:").pack(anchor="w", padx=10)
        customers = ["Wszystkie"] + [c['name'] for c in self.db.get_customers()]
        self.customer_filter = ctk.CTkComboBox(sidebar, values=customers, width=250)
        self.customer_filter.pack(padx=10, pady=5)
        self.customer_filter.set("Wszystkie")
        
        # Status filter
        ctk.CTkLabel(sidebar, text="Status:").pack(anchor="w", padx=10)
        statuses = ["Wszystkie"] + list(STATUS_NAMES.values())
        self.status_filter = ctk.CTkComboBox(sidebar, values=statuses, width=250)
        self.status_filter.pack(padx=10, pady=5)
        self.status_filter.set("Wszystkie")
        
        # Date filters
        ctk.CTkLabel(sidebar, text="Data od:").pack(anchor="w", padx=10, pady=(10, 0))
        self.date_from = DateEntry(sidebar, width=12, background='darkblue',
                                  foreground='white', borderwidth=2)
        self.date_from.pack(padx=10, pady=5)
        
        ctk.CTkLabel(sidebar, text="Data do:").pack(anchor="w", padx=10)
        self.date_to = DateEntry(sidebar, width=12, background='darkblue',
                               foreground='white', borderwidth=2)
        self.date_to.pack(padx=10, pady=5)
        
        # Title search
        ctk.CTkLabel(sidebar, text="Szukaj w tytule:").pack(anchor="w", padx=10, pady=(10, 0))
        self.title_search = ctk.CTkEntry(sidebar, width=250)
        self.title_search.pack(padx=10, pady=5)
        
        # Apply filters button
        ctk.CTkButton(
            sidebar,
            text="Zastosuj filtry",
            width=250,
            command=self.apply_filters
        ).pack(pady=10)
        
        # Dashboard section
        dashboard_label = ctk.CTkLabel(
            sidebar,
            text="Dashboard",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        dashboard_label.pack(pady=(20, 10))
        
        # Status counts frame
        self.status_frame = ctk.CTkFrame(sidebar)
        self.status_frame.pack(padx=10, pady=5, fill="x")
        
        # SLA indicators
        self.sla_frame = ctk.CTkFrame(sidebar)
        self.sla_frame.pack(padx=10, pady=5, fill="x")
        
        # Chart placeholder
        self.chart_frame = ctk.CTkFrame(sidebar, height=200)
        self.chart_frame.pack(padx=10, pady=10, fill="both", expand=True)
    
    def create_orders_area(self, parent):
        """Create orders list area"""
        orders_frame = ctk.CTkFrame(parent)
        orders_frame.grid(row=0, column=1, sticky="nsew")
        
        # Orders list header
        list_header = ctk.CTkFrame(orders_frame)
        list_header.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(
            list_header,
            text="Lista zamówień",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(side="left", padx=10)
        
        # Export buttons
        export_frame = ctk.CTkFrame(list_header)
        export_frame.pack(side="right", padx=10)
        
        ctk.CTkButton(
            export_frame,
            text="📊 Excel",
            width=80,
            command=self.export_excel
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            export_frame,
            text="📄 Word",
            width=80,
            command=self.export_word
        ).pack(side="left", padx=2)
        
        ctk.CTkButton(
            export_frame,
            text="📑 PDF",
            width=80,
            command=self.export_pdf
        ).pack(side="left", padx=2)
        
        # Orders treeview
        from tkinter import ttk
        
        tree_frame = ctk.CTkFrame(orders_frame)
        tree_frame.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Create treeview with scrollbars
        self.orders_tree = ttk.Treeview(
            tree_frame,
            columns=('process_no', 'customer', 'title', 'status', 'price', 
                    'received', 'planned', 'finished'),
            show='headings',
            selectmode='browse'
        )
        
        # Configure columns
        self.orders_tree.heading('process_no', text='Nr procesu')
        self.orders_tree.heading('customer', text='Klient')
        self.orders_tree.heading('title', text='Tytuł')
        self.orders_tree.heading('status', text='Status')
        self.orders_tree.heading('price', text='Cena [PLN]')
        self.orders_tree.heading('received', text='Data wpływu')
        self.orders_tree.heading('planned', text='Planowana')
        self.orders_tree.heading('finished', text='Zakończona')
        
        self.orders_tree.column('process_no', width=100)
        self.orders_tree.column('customer', width=150)
        self.orders_tree.column('title', width=200)
        self.orders_tree.column('status', width=120)
        self.orders_tree.column('price', width=100)
        self.orders_tree.column('received', width=100)
        self.orders_tree.column('planned', width=100)
        self.orders_tree.column('finished', width=100)
        
        # Scrollbars
        v_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.orders_tree.yview)
        h_scroll = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.orders_tree.xview)
        self.orders_tree.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)
        
        # Grid layout
        self.orders_tree.grid(row=0, column=0, sticky="nsew")
        v_scroll.grid(row=0, column=1, sticky="ns")
        h_scroll.grid(row=1, column=0, sticky="ew")
        
        tree_frame.grid_columnconfigure(0, weight=1)
        tree_frame.grid_rowconfigure(0, weight=1)
        
        # Style
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Treeview", background="#212121", foreground="white",
                       fieldbackground="#212121", borderwidth=0)
        style.configure("Treeview.Heading", background="#313131", foreground="white")
        
        # Context menu
        self.create_context_menu()
        
        # Bind events
        self.orders_tree.bind("<Double-1>", self.on_order_double_click)
        self.orders_tree.bind("<Button-3>", self.show_context_menu)
        
        # Apply style with automatic row height for orders list
        settings = get_settings_manager().settings
        configure_treeview_style(self.orders_tree, settings, auto_row_height=True)

        # Define tags for status colors
        for status, color in STATUS_COLORS.items():
            self.orders_tree.tag_configure(status, background=color, foreground='white')
    
    def create_context_menu(self):
        """Create context menu for orders"""
        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="✏️ Edytuj", command=self.edit_order)
        self.context_menu.add_command(label="📁 Pliki", command=self.manage_files)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="📊 Zmień status", command=self.change_status)
        self.context_menu.add_command(label="📜 Historia", command=self.show_history)
        self.context_menu.add_separator()
        self.context_menu.add_command(label="🗑️ Usuń", command=self.delete_order)
    
    def create_status_bar(self):
        """Create status bar"""
        status_bar = ctk.CTkFrame(self, height=30)
        status_bar.grid(row=2, column=0, sticky="ew", padx=5, pady=(0, 5))
        
        self.status_label = ctk.CTkLabel(
            status_bar,
            text=f"Gotowe | Wersja {VERSION} | © 2025",
            font=ctk.CTkFont(size=12)
        )
        self.status_label.pack(side="left", padx=10)
        
        # Connection status
        self.connection_label = ctk.CTkLabel(
            status_bar,
            text="🟢 Połączono",
            font=ctk.CTkFont(size=12)
        )
        self.connection_label.pack(side="right", padx=10)
    
    def load_orders(self, filters=None):
        """Load orders from database"""
        self.orders_tree.delete(*self.orders_tree.get_children())
        
        orders = self.db.get_orders(filters)
        
        today = datetime.date.today()
        
        for order in orders:
            # Format dates
            received = order.get('received_at', '')[:10] if order.get('received_at') else ''
            planned = order.get('planned_at', '')[:10] if order.get('planned_at') else ''
            finished = order.get('finished_at', '')[:10] if order.get('finished_at') else ''
            
            # Format price
            price = f"{order.get('price_pln', 0):.2f}"
            
            # Get status name
            status_name = STATUS_NAMES.get(order.get('status', 'RECEIVED'))
            
            # Determine tag (for coloring based on SLA)
            tag = order.get('status', 'RECEIVED')
            
            # Check if overdue or soon
            if planned and not finished:
                planned_date = datetime.datetime.fromisoformat(order['planned_at']).date()
                days_to_deadline = (planned_date - today).days
                
                if days_to_deadline < 0:
                    tag = 'OVERDUE'
                elif days_to_deadline <= 2:
                    tag = 'SOON'
            
            # Insert into tree
            self.orders_tree.insert('', 'end',
                                   values=(order.get('process_no', ''),
                                         order.get('customer_name', ''),
                                         order.get('title', ''),
                                         status_name,
                                         price,
                                         received,
                                         planned,
                                         finished),
                                   tags=(tag, order['id']))
        
        # Configure special tags
        self.orders_tree.tag_configure('OVERDUE', background='#FF0000', foreground='white')
        self.orders_tree.tag_configure('SOON', background='#FF8C00', foreground='white')
        
        # Update status
        count = len(orders)
        self.status_label.configure(text=f"Załadowano {count} zamówień | Wersja {VERSION}")
    
    def update_dashboard(self):
        """Update dashboard widgets"""
        # Status counts
        counts = self.db.get_status_counts()
        
        # Clear status frame
        for widget in self.status_frame.winfo_children():
            widget.destroy()
        
        ctk.CTkLabel(
            self.status_frame,
            text="Statusy zamówień:",
            font=ctk.CTkFont(weight="bold")
        ).pack(anchor="w", pady=5)
        
        for status, pl_name in STATUS_NAMES.items():
            count = counts.get(status, 0)
            color = STATUS_COLORS[status]
            
            frame = ctk.CTkFrame(self.status_frame)
            frame.pack(fill="x", pady=2)
            
            # Color indicator
            indicator = ctk.CTkLabel(frame, text="●", text_color=color, width=20)
            indicator.pack(side="left")
            
            # Status label
            label = ctk.CTkLabel(frame, text=f"{pl_name}: {count}")
            label.pack(side="left", padx=5)
        
        # SLA indicators
        sla_data = self.db.get_sla_dashboard()
        
        # Clear SLA frame
        for widget in self.sla_frame.winfo_children():
            widget.destroy()
        
        ctk.CTkLabel(
            self.sla_frame,
            text="Terminy realizacji:",
            font=ctk.CTkFont(weight="bold")
        ).pack(anchor="w", pady=5)
        
        # Overdue
        overdue_frame = ctk.CTkFrame(self.sla_frame)
        overdue_frame.pack(fill="x", pady=2)
        ctk.CTkLabel(overdue_frame, text="🔴", width=20).pack(side="left")
        ctk.CTkLabel(overdue_frame, text=f"Przeterminowane: {sla_data['overdue']}").pack(side="left")
        
        # Soon
        soon_frame = ctk.CTkFrame(self.sla_frame)
        soon_frame.pack(fill="x", pady=2)
        ctk.CTkLabel(soon_frame, text="🟡", width=20).pack(side="left")
        ctk.CTkLabel(soon_frame, text=f"Zbliżające się (≤2 dni): {sla_data['soon']}").pack(side="left")
        
        # On time
        on_time_frame = ctk.CTkFrame(self.sla_frame)
        on_time_frame.pack(fill="x", pady=2)
        ctk.CTkLabel(on_time_frame, text="🟢", width=20).pack(side="left")
        ctk.CTkLabel(on_time_frame, text=f"W terminie: {sla_data['on_time']}").pack(side="left")
        
        # Update chart
        self.update_chart(counts)
    
    def update_chart(self, counts):
        """Update status chart"""
        # Clear previous chart
        for widget in self.chart_frame.winfo_children():
            widget.destroy()
        
        if not counts:
            return
        
        # Create pie chart
        fig = Figure(figsize=(3, 2.5), dpi=80)
        ax = fig.add_subplot(111)
        
        labels = [STATUS_NAMES[s] for s in counts.keys()]
        values = list(counts.values())
        colors_list = [STATUS_COLORS[s] for s in counts.keys()]
        
        ax.pie(values, labels=labels, colors=colors_list, autopct='%1.0f%%',
              startangle=90, textprops={'size': 8})
        ax.set_title("Rozkład statusów", fontsize=10)
        
        # Dark theme
        fig.patch.set_facecolor('#212121')
        ax.set_facecolor('#212121')
        
        # Embed in tkinter
        canvas = FigureCanvasTkAgg(fig, master=self.chart_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)
    
    def apply_filters(self):
        """Apply filters to orders list"""
        filters = {}
        
        # Customer filter
        customer = self.customer_filter.get()
        if customer != "Wszystkie":
            customer_id = None
            for c in self.db.get_customers():
                if c['name'] == customer:
                    customer_id = c['id']
                    break
            if customer_id:
                filters['customer_id'] = customer_id
        
        # Status filter
        status = self.status_filter.get()
        if status != "Wszystkie":
            for key, value in STATUS_NAMES.items():
                if value == status:
                    filters['status'] = key
                    break
        
        # Date filters
        date_from = self.date_from.get_date()
        date_to = self.date_to.get_date()
        if date_from:
            filters['date_from'] = date_from.isoformat()
        if date_to:
            filters['date_to'] = date_to.isoformat()
        
        # Title search
        title = self.title_search.get().strip()
        if title:
            filters['title'] = title
        
        self.load_orders(filters)
    
    def new_order(self):
        """Create new order"""
        dialog = OrderDialog(self, self.db)
        self.wait_window(dialog)
        self.refresh_all()

    def generate_wz(self):
        """Generuje dokument WZ dla wybranego zamówienia"""
        # Pobierz wybrane zamówienie z listy
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie do wygenerowania WZ")
            return

        # Pobierz ID zamówienia z tagów
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]  # Drugi tag to ID zamówienia

        # Otwórz dialog generowania WZ
        try:
            dialog = WZGeneratorDialog(self, self.db.client, order_id)
            dialog.focus()
        except Exception as e:
            messagebox.showerror("Błąd", f"Nie można otworzyć dialogu WZ:\n{e}")

    def edit_order(self):
        """Edit selected order"""
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie do edycji")
            return
        
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]
        
        # Get full order data
        orders = self.db.get_orders()
        order_data = None
        for order in orders:
            if order['id'] == order_id:
                order_data = order
                break
        
        if order_data:
            dialog = OrderDialog(self, self.db, order_data)
            self.wait_window(dialog)
            self.refresh_all()
    
    def delete_order(self):
        """Delete selected order"""
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie do usunięcia")
            return
        
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]
        process_no = item['values'][0]
        
        if messagebox.askyesno("Potwierdzenie",
                               f"Czy na pewno usunąć zamówienie {process_no}?"):
            if self.db.delete_order(order_id):
                messagebox.showinfo("Sukces", "Zamówienie zostało usunięte")
                self.refresh_all()
            else:
                messagebox.showerror("Błąd", "Nie udało się usunąć zamówienia")
    
    def manage_customers(self):
        """Open customers management dialog"""
        dialog = CustomerDialog(self, self.db)
        self.wait_window(dialog)

        # Refresh customer filter
        customers = ["Wszystkie"] + [c['name'] for c in self.db.get_customers()]
        self.customer_filter.configure(values=customers)

    def open_settings(self):
        """Open settings dialog"""
        dialog = SettingsDialog(self, callback=self.on_settings_changed)
        self.wait_window(dialog)

    def on_settings_changed(self):
        """Callback when settings are changed"""
        # Reload settings
        self.settings = self.settings_manager.settings

        # Update logo if changed
        self.update_logo_display()

        # Refresh UI if needed
        if self.settings.list_show_thumbnails:
            self.load_orders()

    def update_logo_display(self):
        """Update logo display in header"""
        # This will be called when settings change
        # The logo will be updated on next app restart
        pass

    def manage_files(self):
        """Manage files for selected order"""
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie")
            return
        
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]
        process_no = item['values'][0]
        
        # File dialog
        file_path = filedialog.askopenfilename(
            title=f"Wybierz plik dla zamówienia {process_no}",
            filetypes=[
                ("Wszystkie pliki", "*.*"),
                ("DXF", "*.dxf"),
                ("DWG", "*.dwg"),
                ("STP", "*.stp"),
                ("PDF", "*.pdf"),
                ("Obrazy", "*.png *.jpg *.jpeg *.gif")
            ]
        )
        
        if file_path:
            if self.db.upload_file(order_id, process_no, file_path):
                messagebox.showinfo("Sukces", "Plik został przesłany")
            else:
                messagebox.showerror("Błąd", "Nie udało się przesłać pliku")
    
    def change_status(self):
        """Change status of selected order"""
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie")
            return
        
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]
        current_status_pl = item['values'][3]
        
        # Status change dialog
        dialog = ctk.CTkToplevel(self)
        dialog.title("Zmiana statusu")
        dialog.geometry("400x200")
        dialog.transient(self)
        dialog.grab_set()
        
        # Center
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (400 // 2)
        y = (dialog.winfo_screenheight() // 2) - (200 // 2)
        dialog.geometry(f"+{x}+{y}")
        
        # Content
        frame = ctk.CTkFrame(dialog)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        ctk.CTkLabel(
            frame,
            text=f"Obecny status: {current_status_pl}",
            font=ctk.CTkFont(size=14)
        ).pack(pady=10)
        
        ctk.CTkLabel(frame, text="Nowy status:").pack(pady=5)
        
        status_combo = ctk.CTkComboBox(
            frame,
            values=list(STATUS_NAMES.values()),
            width=250
        )
        status_combo.pack(pady=5)
        status_combo.set(current_status_pl)
        
        def save_status():
            new_status_pl = status_combo.get()
            new_status = None
            for key, value in STATUS_NAMES.items():
                if value == new_status_pl:
                    new_status = key
                    break
            
            if new_status:
                updates = {'status': new_status}
                if new_status == 'DONE' or new_status == 'INVOICED':
                    updates['finished_at'] = datetime.date.today().isoformat()
                
                if self.db.update_order(order_id, updates):
                    messagebox.showinfo("Sukces", "Status został zmieniony")
                    dialog.destroy()
                    self.refresh_all()
                else:
                    messagebox.showerror("Błąd", "Nie udało się zmienić statusu")
        
        ctk.CTkButton(
            frame,
            text="Zapisz",
            command=save_status
        ).pack(pady=20)
    
    def show_history(self):
        """Show status history for selected order"""
        selection = self.orders_tree.selection()
        if not selection:
            messagebox.showwarning("Uwaga", "Wybierz zamówienie")
            return
        
        item = self.orders_tree.item(selection[0])
        order_id = item['tags'][1]
        process_no = item['values'][0]
        
        # Get history
        try:
            response = self.db.client.table('order_status_history').select("*") \
                       .eq('order_id', order_id).order('changed_at', desc=True).execute()
            history = response.data
        except:
            history = []
        
        # History dialog
        dialog = ctk.CTkToplevel(self)
        dialog.title(f"Historia statusów - {process_no}")
        dialog.geometry("600x400")
        dialog.transient(self)
        dialog.grab_set()
        
        # Center
        dialog.update_idletasks()
        x = (dialog.winfo_screenwidth() // 2) - (600 // 2)
        y = (dialog.winfo_screenheight() // 2) - (400 // 2)
        dialog.geometry(f"+{x}+{y}")
        
        # Content
        frame = ctk.CTkScrollableFrame(dialog)
        frame.pack(fill="both", expand=True, padx=20, pady=20)
        
        if history:
            for entry in history:
                date = entry.get('changed_at', '')[:19].replace('T', ' ')
                old_status = STATUS_NAMES.get(entry.get('old_status', ''), entry.get('old_status', ''))
                new_status = STATUS_NAMES.get(entry.get('new_status', ''), entry.get('new_status', ''))
                
                entry_frame = ctk.CTkFrame(frame)
                entry_frame.pack(fill="x", pady=5)
                
                ctk.CTkLabel(
                    entry_frame,
                    text=f"📅 {date}",
                    font=ctk.CTkFont(weight="bold")
                ).pack(anchor="w", padx=10, pady=2)
                
                ctk.CTkLabel(
                    entry_frame,
                    text=f"{old_status} ➡️ {new_status}"
                ).pack(anchor="w", padx=20)
        else:
            ctk.CTkLabel(
                frame,
                text="Brak historii zmian statusu"
            ).pack(pady=20)
    
    def show_reports(self):
        """Show reports dialog"""
        dialog = ReportsDialog(self, self.db)
    
    def export_excel(self):
        """Export orders to Excel"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")]
        )
        
        if not file_path:
            return
        
        try:
            # Get visible orders
            orders = []
            for child in self.orders_tree.get_children():
                item = self.orders_tree.item(child)
                orders.append(item['values'])
            
            # Create DataFrame
            df = pd.DataFrame(orders, columns=[
                'Nr procesu', 'Klient', 'Tytuł', 'Status', 
                'Cena [PLN]', 'Data wpływu', 'Planowana', 'Zakończona'
            ])
            
            # Save to Excel with formatting
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Zamówienia', index=False)
                
                # Get workbook and worksheet
                workbook = writer.book
                worksheet = writer.sheets['Zamówienia']
                
                # Format header
                for cell in worksheet[1]:
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="366092")
                    cell.alignment = Alignment(horizontal="center")
                
                # Auto-fit columns
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    worksheet.column_dimensions[column_letter].width = adjusted_width
            
            messagebox.showinfo("Sukces", f"Eksport zakończony:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Błąd", f"Błąd eksportu:\n{e}")
    
    def export_word(self):
        """Export orders to Word"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")]
        )
        
        if not file_path:
            return
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Raport zamówień - Laser/Prasa', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            doc.add_paragraph(f'Data raportu: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
            
            # Table
            table = doc.add_table(rows=1, cols=8)
            table.style = 'Light Grid Accent 1'
            
            # Header
            headers = ['Nr procesu', 'Klient', 'Tytuł', 'Status', 
                      'Cena [PLN]', 'Data wpływu', 'Planowana', 'Zakończona']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data
            for child in self.orders_tree.get_children():
                item = self.orders_tree.item(child)
                values = item['values']
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
            
            doc.save(file_path)
            messagebox.showinfo("Sukces", f"Eksport zakończony:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Błąd", f"Błąd eksportu:\n{e}")
    
    def export_pdf(self):
        """Export orders to PDF"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")]
        )
        
        if not file_path:
            return
        
        try:
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            elements = []
            
            # Title
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1F4E79'),
                spaceAfter=30,
                alignment=1  # Center
            )
            
            elements.append(Paragraph("Raport zamówień - Laser/Prasa", title_style))
            elements.append(Spacer(1, 12))
            
            # Date
            elements.append(Paragraph(f"Data raportu: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}", 
                                    styles['Normal']))
            elements.append(Spacer(1, 20))
            
            # Table data
            data = [['Nr procesu', 'Klient', 'Tytuł', 'Status', 'Cena [PLN]']]
            
            for child in self.orders_tree.get_children():
                item = self.orders_tree.item(child)
                values = item['values']
                data.append([
                    values[0],  # process_no
                    values[1],  # customer
                    values[2][:30] + '...' if len(str(values[2])) > 30 else values[2],  # title (truncated)
                    values[3],  # status
                    values[4]   # price
                ])
            
            # Create table
            table = Table(data)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(table)
            
            # Build PDF
            doc.build(elements)
            messagebox.showinfo("Sukces", f"Eksport zakończony:\n{file_path}")
            
        except Exception as e:
            messagebox.showerror("Błąd", f"Błąd eksportu:\n{e}")
    
    def on_order_double_click(self, event):
        """Handle double click on order"""
        self.edit_order()
    
    def show_context_menu(self, event):
        """Show context menu"""
        # Select item under cursor
        item = self.orders_tree.identify('item', event.x, event.y)
        if item:
            self.orders_tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)
    
    def refresh_all(self):
        """Refresh all data"""
        self.load_orders()
        self.update_dashboard()
        self.status_label.configure(text="Odświeżono dane")
    
    def auto_refresh(self):
        """Auto refresh every 5 minutes"""
        self.refresh_all()
        self.after(300000, self.auto_refresh)  # 5 minutes

class ReportsDialog(ctk.CTkToplevel):
    """Advanced reports dialog"""
    
    def __init__(self, parent, db: SupabaseManager):
        super().__init__(parent)
        self.db = db
        
        self.title("Raporty i analizy")
        self.geometry("1200x700")
        
        self.setup_ui()
        
        # Center window
        self.update_idletasks()
        x = (self.winfo_screenwidth() // 2) - (1200 // 2)
        y = (self.winfo_screenheight() // 2) - (700 // 2)
        self.geometry(f"+{x}+{y}")
    
    def setup_ui(self):
        """Setup reports UI"""
        # Main container
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Header
        header_frame = ctk.CTkFrame(main_frame)
        header_frame.pack(fill="x", padx=5, pady=5)
        
        ctk.CTkLabel(
            header_frame,
            text="Raporty i analizy",
            font=ctk.CTkFont(size=20, weight="bold")
        ).pack(side="left", padx=10)
        
        # Tab view
        self.tabview = ctk.CTkTabview(main_frame)
        self.tabview.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Create tabs
        self.tabview.add("📊 Wykresy")
        self.tabview.add("📈 Statystyki")
        self.tabview.add("💰 Finansowe")
        self.tabview.add("⏰ Terminy")
        
        # Setup each tab
        self.setup_charts_tab()
        self.setup_statistics_tab()
        self.setup_financial_tab()
        self.setup_deadlines_tab()
    
    def setup_charts_tab(self):
        """Setup charts tab"""
        tab = self.tabview.tab("📊 Wykresy")
        
        # Create matplotlib figure with subplots
        fig = Figure(figsize=(12, 6), dpi=80)
        
        # Orders by month
        ax1 = fig.add_subplot(2, 2, 1)
        orders = self.db.get_orders()
        
        # Group by month
        months = {}
        for order in orders:
            if order.get('received_at'):
                month = order['received_at'][:7]  # YYYY-MM
                months[month] = months.get(month, 0) + 1
        
        if months:
            ax1.bar(months.keys(), months.values(), color='#4169E1')
            ax1.set_title('Zamówienia wg miesięcy')
            ax1.set_xlabel('Miesiąc')
            ax1.set_ylabel('Liczba')
            ax1.tick_params(axis='x', rotation=45)
        
        # Orders by customer (top 10)
        ax2 = fig.add_subplot(2, 2, 2)
        customers = {}
        for order in orders:
            customer = order.get('customer_name', 'Brak')
            customers[customer] = customers.get(customer, 0) + 1
        
        if customers:
            sorted_customers = dict(sorted(customers.items(), key=lambda x: x[1], reverse=True)[:10])
            ax2.barh(list(sorted_customers.keys()), list(sorted_customers.values()), color='#32CD32')
            ax2.set_title('Top 10 klientów')
            ax2.set_xlabel('Liczba zamówień')
        
        # Revenue by month
        ax3 = fig.add_subplot(2, 2, 3)
        revenue = {}
        for order in orders:
            if order.get('received_at'):
                month = order['received_at'][:7]
                price = float(order.get('price_pln', 0) or 0)
                revenue[month] = revenue.get(month, 0) + price
        
        if revenue:
            ax3.plot(list(revenue.keys()), list(revenue.values()), marker='o', color='#FFD700')
            ax3.set_title('Przychód wg miesięcy')
            ax3.set_xlabel('Miesiąc')
            ax3.set_ylabel('PLN')
            ax3.tick_params(axis='x', rotation=45)
        
        # Status distribution over time
        ax4 = fig.add_subplot(2, 2, 4)
        status_counts = self.db.get_status_counts()
        if status_counts:
            labels = [STATUS_NAMES[s] for s in status_counts.keys()]
            sizes = list(status_counts.values())
            colors_list = [STATUS_COLORS[s] for s in status_counts.keys()]
            ax4.pie(sizes, labels=labels, colors=colors_list, autopct='%1.1f%%')
            ax4.set_title('Rozkład statusów')
        
        # Dark theme
        fig.patch.set_facecolor('#212121')
        for ax in [ax1, ax2, ax3, ax4]:
            ax.set_facecolor('#313131')
            ax.tick_params(colors='white')
            ax.title.set_color('white')
            ax.xaxis.label.set_color('white')
            ax.yaxis.label.set_color('white')
            for spine in ax.spines.values():
                spine.set_edgecolor('white')
        
        fig.tight_layout()
        
        # Embed in tkinter
        canvas = FigureCanvasTkAgg(fig, master=tab)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)
    
    def setup_statistics_tab(self):
        """Setup statistics tab"""
        tab = self.tabview.tab("📈 Statystyki")
        
        # Create scrollable frame
        scroll_frame = ctk.CTkScrollableFrame(tab)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        orders = self.db.get_orders()
        
        # General statistics
        stats_frame = ctk.CTkFrame(scroll_frame)
        stats_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            stats_frame,
            text="Statystyki ogólne",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(anchor="w", padx=10, pady=5)
        
        # Calculate statistics
        total_orders = len(orders)
        total_revenue = sum(float(o.get('price_pln', 0) or 0) for o in orders)
        avg_price = total_revenue / total_orders if total_orders > 0 else 0
        
        # Completion rate
        completed = len([o for o in orders if o.get('status') in ['DONE', 'INVOICED']])
        completion_rate = (completed / total_orders * 100) if total_orders > 0 else 0
        
        # Average processing time
        processing_times = []
        for order in orders:
            if order.get('received_at') and order.get('finished_at'):
                start = datetime.datetime.fromisoformat(order['received_at'])
                end = datetime.datetime.fromisoformat(order['finished_at'])
                processing_times.append((end - start).days)
        
        avg_processing = sum(processing_times) / len(processing_times) if processing_times else 0
        
        # Display statistics
        stats_grid = ctk.CTkFrame(stats_frame)
        stats_grid.pack(padx=20, pady=10)
        
        stats = [
            ("Łączna liczba zamówień:", f"{total_orders}"),
            ("Łączny przychód:", f"{total_revenue:.2f} PLN"),
            ("Średnia wartość zamówienia:", f"{avg_price:.2f} PLN"),
            ("Wskaźnik realizacji:", f"{completion_rate:.1f}%"),
            ("Średni czas realizacji:", f"{avg_processing:.1f} dni"),
            ("Liczba klientów:", f"{len(set(o.get('customer_id') for o in orders))}"),
        ]
        
        for i, (label, value) in enumerate(stats):
            ctk.CTkLabel(stats_grid, text=label, font=ctk.CTkFont(size=14)).grid(row=i, column=0, padx=10, pady=5, sticky="w")
            ctk.CTkLabel(stats_grid, text=value, font=ctk.CTkFont(size=14, weight="bold")).grid(row=i, column=1, padx=10, pady=5, sticky="w")
    
    def setup_financial_tab(self):
        """Setup financial tab"""
        tab = self.tabview.tab("💰 Finansowe")
        
        scroll_frame = ctk.CTkScrollableFrame(tab)
        scroll_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        orders = self.db.get_orders()
        
        # Financial summary
        summary_frame = ctk.CTkFrame(scroll_frame)
        summary_frame.pack(fill="x", pady=10)
        
        ctk.CTkLabel(
            summary_frame,
            text="Podsumowanie finansowe",
            font=ctk.CTkFont(size=18, weight="bold")
        ).pack(anchor="w", padx=10, pady=5)
        
        # Calculate by status
        by_status = {}
        for order in orders:
            status = order.get('status', 'RECEIVED')
            price = float(order.get('price_pln', 0) or 0)
            if status not in by_status:
                by_status[status] = {'count': 0, 'total': 0}
            by_status[status]['count'] += 1
            by_status[status]['total'] += price
        
        # Display
        for status, data in by_status.items():
            status_name = STATUS_NAMES.get(status, status)
            frame = ctk.CTkFrame(summary_frame)
            frame.pack(fill="x", padx=20, pady=5)
            
            ctk.CTkLabel(
                frame,
                text=f"{status_name}:",
                font=ctk.CTkFont(size=14)
            ).pack(side="left", padx=10)
            
            ctk.CTkLabel(
                frame,
                text=f"{data['count']} zamówień",
                font=ctk.CTkFont(size=14)
            ).pack(side="left", padx=10)
            
            ctk.CTkLabel(
                frame,
                text=f"{data['total']:.2f} PLN",
                font=ctk.CTkFont(size=14, weight="bold")
            ).pack(side="left", padx=10)
    
    def setup_deadlines_tab(self):
        """Setup deadlines tab"""
        tab = self.tabview.tab("⏰ Terminy")
        
        # Get SLA data
        sla_orders = []
        orders = self.db.get_orders()
        today = datetime.date.today()
        
        for order in orders:
            if order.get('planned_at') and not order.get('finished_at'):
                planned_date = datetime.datetime.fromisoformat(order['planned_at']).date()
                days_to_deadline = (planned_date - today).days
                
                sla_orders.append({
                    'process_no': order.get('process_no', ''),
                    'customer': order.get('customer_name', ''),
                    'title': order.get('title', ''),
                    'planned': planned_date,
                    'days': days_to_deadline,
                    'status': 'Przeterminowane' if days_to_deadline < 0 else 
                             'Pilne' if days_to_deadline <= 2 else 'W terminie'
                })
        
        # Sort by days to deadline
        sla_orders.sort(key=lambda x: x['days'])
        
        # Create treeview
        from tkinter import ttk
        
        tree = ttk.Treeview(
            tab,
            columns=('process_no', 'customer', 'title', 'planned', 'days', 'status'),
            show='headings',
            height=15
        )
        
        tree.heading('process_no', text='Nr procesu')
        tree.heading('customer', text='Klient')
        tree.heading('title', text='Tytuł')
        tree.heading('planned', text='Planowana data')
        tree.heading('days', text='Dni do terminu')
        tree.heading('status', text='Status SLA')
        
        tree.column('process_no', width=100)
        tree.column('customer', width=150)
        tree.column('title', width=250)
        tree.column('planned', width=100)
        tree.column('days', width=100)
        tree.column('status', width=120)
        
        # Insert data with colors
        for order in sla_orders:
            tag = 'overdue' if order['days'] < 0 else 'urgent' if order['days'] <= 2 else 'ok'
            tree.insert('', 'end', values=(
                order['process_no'],
                order['customer'],
                order['title'],
                order['planned'],
                order['days'],
                order['status']
            ), tags=(tag,))
        
        # Configure tags
        tree.tag_configure('overdue', background='#FF0000', foreground='white')
        tree.tag_configure('urgent', background='#FF8C00', foreground='white')
        tree.tag_configure('ok', background='#212121', foreground='white')
        
        tree.pack(fill="both", expand=True, padx=10, pady=10)

def main():
    """Main entry point"""
    try:
        app = MainApplication()
        app.mainloop()
    except Exception as e:
        messagebox.showerror("Błąd krytyczny", f"Wystąpił błąd:\n{e}")
        sys.exit(1)

if __name__ == "__main__":
    main()

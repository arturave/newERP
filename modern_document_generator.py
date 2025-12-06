#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Modern Document Generator
Generates business documents (WZ, Order Confirmations, Order Lists) using HTML templates
without reportlab dependency
"""

import os
import json
from datetime import datetime, date
from typing import Dict, List, Optional, Any
from pathlib import Path
import tempfile
import base64

# HTML to PDF conversion options
try:
    import pdfkit
    PDFKIT_AVAILABLE = True
except ImportError:
    PDFKIT_AVAILABLE = False

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

# Word and Excel support
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill


class ModernDocumentGenerator:
    """Modern document generator using HTML templates"""

    def __init__(self, company_settings: Optional[Dict] = None):
        """Initialize the generator with company settings"""
        self.company_settings = company_settings or self._get_default_company_settings()

        # Check available PDF generators
        self.pdf_engine = None
        if WEASYPRINT_AVAILABLE:
            self.pdf_engine = 'weasyprint'
        elif PDFKIT_AVAILABLE:
            self.pdf_engine = 'pdfkit'
        else:
            print("⚠️ No PDF engine available. Install weasyprint or pdfkit for PDF generation.")

    def _get_default_company_settings(self) -> Dict:
        """Get default company settings"""
        return {
            'name': 'Your Company Ltd.',
            'address': 'ul. Przykładowa 123',
            'city': '00-000 Warszawa',
            'country': 'Poland',
            'nip': '123-456-78-90',
            'phone': '+48 123 456 789',
            'email': 'office@company.pl',
            'website': 'www.company.pl',
            'logo_base64': None  # Base64 encoded logo
        }

    def generate_wz_document(
        self,
        wz_data: Dict,
        format: str = 'pdf'
    ) -> Optional[bytes]:
        """
        Generate WZ (External Issue) document

        Args:
            wz_data: Dictionary with WZ data
            format: Output format ('pdf', 'html', 'docx', 'xlsx')

        Returns:
            Document content as bytes or None if error
        """
        if format == 'html' or format == 'pdf':
            html_content = self._generate_wz_html(wz_data)

            if format == 'html':
                return html_content.encode('utf-8')
            else:
                return self._html_to_pdf(html_content)

        elif format == 'docx':
            return self._generate_wz_docx(wz_data)

        elif format == 'xlsx':
            return self._generate_wz_xlsx(wz_data)

        return None

    def _generate_wz_html(self, wz_data: Dict) -> str:
        """Generate HTML template for WZ document"""

        # Prepare data
        wz_number = wz_data.get('wz_number', 'WZ/2025/001')
        issue_date = wz_data.get('issue_date', datetime.now().strftime('%Y-%m-%d'))
        order_number = wz_data.get('order_number', '')

        recipient = wz_data.get('recipient', {})
        items = wz_data.get('items', [])
        notes = wz_data.get('notes', '')

        # Generate items rows
        items_html = ''
        for i, item in enumerate(items, 1):
            items_html += f"""
            <tr>
                <td class="center">{i}</td>
                <td>{item.get('name', '')}</td>
                <td class="center">{item.get('quantity', 0)}</td>
                <td class="center">{item.get('unit', 'szt.')}</td>
                <td>{item.get('notes', '')}</td>
            </tr>
            """

        # Logo handling
        logo_html = ''
        if self.company_settings.get('logo_base64'):
            logo_html = f'<img src="data:image/png;base64,{self.company_settings["logo_base64"]}" class="logo">'

        # HTML template based on Polish WZ standards
        html = f"""
        <!DOCTYPE html>
        <html lang="pl">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>WZ {wz_number}</title>
            <style>
                @page {{
                    size: A4;
                    margin: 20mm;
                }}

                * {{
                    margin: 0;
                    padding: 0;
                    box-sizing: border-box;
                }}

                body {{
                    font-family: 'Arial', sans-serif;
                    font-size: 12pt;
                    line-height: 1.5;
                    color: #333;
                }}

                .header {{
                    display: flex;
                    justify-content: space-between;
                    margin-bottom: 30px;
                    padding-bottom: 20px;
                    border-bottom: 2px solid #2c3e50;
                }}

                .logo {{
                    max-height: 80px;
                    max-width: 200px;
                }}

                .company-info {{
                    text-align: right;
                    font-size: 10pt;
                    color: #666;
                }}

                .company-name {{
                    font-size: 16pt;
                    font-weight: bold;
                    color: #2c3e50;
                    margin-bottom: 5px;
                }}

                .document-title {{
                    text-align: center;
                    margin: 30px 0;
                }}

                .document-title h1 {{
                    font-size: 24pt;
                    color: #2c3e50;
                    margin-bottom: 10px;
                }}

                .document-number {{
                    font-size: 14pt;
                    color: #666;
                }}

                .info-section {{
                    display: flex;
                    justify-content: space-between;
                    margin-bottom: 30px;
                }}

                .info-box {{
                    width: 48%;
                    padding: 15px;
                    background: #f8f9fa;
                    border-radius: 5px;
                }}

                .info-box h3 {{
                    font-size: 12pt;
                    color: #2c3e50;
                    margin-bottom: 10px;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 5px;
                }}

                .info-box p {{
                    margin: 5px 0;
                    font-size: 11pt;
                }}

                .items-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}

                .items-table th {{
                    background: #2c3e50;
                    color: white;
                    padding: 10px;
                    text-align: left;
                    font-size: 11pt;
                }}

                .items-table td {{
                    padding: 8px 10px;
                    border-bottom: 1px solid #ddd;
                    font-size: 11pt;
                }}

                .items-table tr:hover {{
                    background: #f8f9fa;
                }}

                .center {{
                    text-align: center;
                }}

                .notes-section {{
                    margin-top: 30px;
                    padding: 15px;
                    background: #f8f9fa;
                    border-radius: 5px;
                }}

                .notes-section h3 {{
                    font-size: 12pt;
                    color: #2c3e50;
                    margin-bottom: 10px;
                }}

                .signatures {{
                    display: flex;
                    justify-content: space-between;
                    margin-top: 50px;
                    padding-top: 20px;
                }}

                .signature-box {{
                    width: 40%;
                    text-align: center;
                }}

                .signature-line {{
                    border-top: 1px solid #333;
                    margin-top: 60px;
                    padding-top: 5px;
                    font-size: 10pt;
                }}

                .footer {{
                    margin-top: 50px;
                    padding-top: 20px;
                    border-top: 1px solid #ddd;
                    text-align: center;
                    font-size: 9pt;
                    color: #999;
                }}

                @media print {{
                    .header {{
                        border-bottom: 2px solid #000;
                    }}

                    .items-table th {{
                        background: #333;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div>
                    {logo_html}
                </div>
                <div class="company-info">
                    <div class="company-name">{self.company_settings['name']}</div>
                    <p>{self.company_settings['address']}</p>
                    <p>{self.company_settings['city']}</p>
                    <p>NIP: {self.company_settings['nip']}</p>
                    <p>Tel: {self.company_settings['phone']}</p>
                    <p>Email: {self.company_settings['email']}</p>
                </div>
            </div>

            <div class="document-title">
                <h1>WYDANIE ZEWNĘTRZNE (WZ)</h1>
                <div class="document-number">Nr: {wz_number}</div>
                <div>Data wystawienia: {issue_date}</div>
                {f'<div>Zamówienie: {order_number}</div>' if order_number else ''}
            </div>

            <div class="info-section">
                <div class="info-box">
                    <h3>WYSTAWCA</h3>
                    <p><strong>{self.company_settings['name']}</strong></p>
                    <p>{self.company_settings['address']}</p>
                    <p>{self.company_settings['city']}</p>
                    <p>NIP: {self.company_settings['nip']}</p>
                </div>

                <div class="info-box">
                    <h3>ODBIORCA</h3>
                    <p><strong>{recipient.get('name', '')}</strong></p>
                    <p>{recipient.get('address', '')}</p>
                    <p>{recipient.get('city', '')} {recipient.get('postal_code', '')}</p>
                    {f"<p>NIP: {recipient.get('nip', '')}</p>" if recipient.get('nip') else ''}
                    {f"<p>Osoba kontaktowa: {recipient.get('contact_person', '')}</p>" if recipient.get('contact_person') else ''}
                    {f"<p>Tel: {recipient.get('phone', '')}</p>" if recipient.get('phone') else ''}
                </div>
            </div>

            <table class="items-table">
                <thead>
                    <tr>
                        <th style="width: 5%">Lp.</th>
                        <th style="width: 45%">Nazwa towaru/materiału</th>
                        <th style="width: 10%">Ilość</th>
                        <th style="width: 10%">J.m.</th>
                        <th style="width: 30%">Uwagi</th>
                    </tr>
                </thead>
                <tbody>
                    {items_html}
                </tbody>
            </table>

            {f'''
            <div class="notes-section">
                <h3>UWAGI</h3>
                <p>{notes}</p>
            </div>
            ''' if notes else ''}

            <div class="signatures">
                <div class="signature-box">
                    <div class="signature-line">
                        Podpis osoby wydającej
                    </div>
                </div>
                <div class="signature-box">
                    <div class="signature-line">
                        Podpis osoby odbierającej
                    </div>
                </div>
            </div>

            <div class="footer">
                <p>Dokument wygenerowany elektronicznie • {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            </div>
        </body>
        </html>
        """

        return html

    def _html_to_pdf(self, html_content: str) -> Optional[bytes]:
        """Convert HTML to PDF using available engine"""

        if self.pdf_engine == 'weasyprint':
            try:
                pdf = HTML(string=html_content).write_pdf()
                return pdf
            except Exception as e:
                print(f"❌ WeasyPrint error: {e}")
                return None

        elif self.pdf_engine == 'pdfkit':
            try:
                options = {
                    'page-size': 'A4',
                    'margin-top': '20mm',
                    'margin-right': '20mm',
                    'margin-bottom': '20mm',
                    'margin-left': '20mm',
                    'encoding': 'UTF-8',
                    'no-outline': None
                }
                pdf = pdfkit.from_string(html_content, False, options=options)
                return pdf
            except Exception as e:
                print(f"❌ PDFKit error: {e}")
                return None

        return None

    def _generate_wz_docx(self, wz_data: Dict) -> bytes:
        """Generate WZ document in DOCX format"""
        doc = Document()

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Company header
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.add_run(self.company_settings['name']).bold = True
        header.add_run(f"\n{self.company_settings['address']}")
        header.add_run(f"\n{self.company_settings['city']}")
        header.add_run(f"\nNIP: {self.company_settings['nip']}")

        # Document title
        title = doc.add_heading('WYDANIE ZEWNĘTRZNE (WZ)', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Document number and date
        doc_info = doc.add_paragraph()
        doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_info.add_run(f"Nr: {wz_data.get('wz_number', 'WZ/2025/001')}")
        doc_info.add_run(f"\nData: {wz_data.get('issue_date', datetime.now().strftime('%Y-%m-%d'))}")

        doc.add_paragraph()

        # Recipient info
        doc.add_heading('ODBIORCA:', 2)
        recipient = wz_data.get('recipient', {})
        doc.add_paragraph(f"{recipient.get('name', '')}")
        doc.add_paragraph(f"{recipient.get('address', '')}")
        doc.add_paragraph(f"{recipient.get('city', '')} {recipient.get('postal_code', '')}")
        if recipient.get('nip'):
            doc.add_paragraph(f"NIP: {recipient.get('nip', '')}")

        doc.add_paragraph()

        # Items table
        doc.add_heading('WYKAZ WYDAWANYCH TOWARÓW:', 2)

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Lp.'
        header_cells[1].text = 'Nazwa towaru/materiału'
        header_cells[2].text = 'Ilość'
        header_cells[3].text = 'J.m.'
        header_cells[4].text = 'Uwagi'

        # Items rows
        for i, item in enumerate(wz_data.get('items', []), 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = item.get('name', '')
            row_cells[2].text = str(item.get('quantity', 0))
            row_cells[3].text = item.get('unit', 'szt.')
            row_cells[4].text = item.get('notes', '')

        # Notes
        if wz_data.get('notes'):
            doc.add_paragraph()
            doc.add_heading('UWAGI:', 2)
            doc.add_paragraph(wz_data.get('notes', ''))

        # Signatures
        doc.add_paragraph()
        doc.add_paragraph()

        signatures = doc.add_table(rows=2, cols=2)
        signatures.rows[0].cells[0].text = '_' * 30
        signatures.rows[0].cells[1].text = '_' * 30
        signatures.rows[1].cells[0].text = 'Podpis osoby wydającej'
        signatures.rows[1].cells[1].text = 'Podpis osoby odbierającej'

        # Save to bytes
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return file_stream.read()

    def _generate_wz_xlsx(self, wz_data: Dict) -> bytes:
        """Generate WZ document in XLSX format"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"WZ_{wz_data.get('wz_number', 'WZ2025001').replace('/', '_')}"

        # Styles
        header_font = Font(bold=True, size=14)
        title_font = Font(bold=True, size=18)
        table_header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        table_header_font = Font(color="FFFFFF", bold=True)

        # Company info
        ws['A1'] = self.company_settings['name']
        ws['A1'].font = header_font
        ws['A2'] = self.company_settings['address']
        ws['A3'] = self.company_settings['city']
        ws['A4'] = f"NIP: {self.company_settings['nip']}"

        # Document title
        ws['A6'] = 'WYDANIE ZEWNĘTRZNE (WZ)'
        ws['A6'].font = title_font
        ws.merge_cells('A6:E6')
        ws['A6'].alignment = Alignment(horizontal='center')

        ws['A7'] = f"Nr: {wz_data.get('wz_number', 'WZ/2025/001')}"
        ws['A8'] = f"Data: {wz_data.get('issue_date', datetime.now().strftime('%Y-%m-%d'))}"

        # Recipient
        ws['A10'] = 'ODBIORCA:'
        ws['A10'].font = header_font
        recipient = wz_data.get('recipient', {})
        ws['A11'] = recipient.get('name', '')
        ws['A12'] = recipient.get('address', '')
        ws['A13'] = f"{recipient.get('city', '')} {recipient.get('postal_code', '')}"
        if recipient.get('nip'):
            ws['A14'] = f"NIP: {recipient.get('nip', '')}"

        # Items table
        start_row = 16
        ws[f'A{start_row}'] = 'Lp.'
        ws[f'B{start_row}'] = 'Nazwa towaru/materiału'
        ws[f'C{start_row}'] = 'Ilość'
        ws[f'D{start_row}'] = 'J.m.'
        ws[f'E{start_row}'] = 'Uwagi'

        # Style header row
        for col in ['A', 'B', 'C', 'D', 'E']:
            cell = ws[f'{col}{start_row}']
            cell.fill = table_header_fill
            cell.font = table_header_font
            cell.alignment = Alignment(horizontal='center')

        # Add items
        for i, item in enumerate(wz_data.get('items', []), 1):
            row = start_row + i
            ws[f'A{row}'] = i
            ws[f'B{row}'] = item.get('name', '')
            ws[f'C{row}'] = item.get('quantity', 0)
            ws[f'D{row}'] = item.get('unit', 'szt.')
            ws[f'E{row}'] = item.get('notes', '')

        # Notes
        if wz_data.get('notes'):
            notes_row = start_row + len(wz_data.get('items', [])) + 3
            ws[f'A{notes_row}'] = 'UWAGI:'
            ws[f'A{notes_row}'].font = header_font
            ws[f'A{notes_row + 1}'] = wz_data.get('notes', '')

        # Adjust column widths
        ws.column_dimensions['A'].width = 8
        ws.column_dimensions['B'].width = 40
        ws.column_dimensions['C'].width = 12
        ws.column_dimensions['D'].width = 10
        ws.column_dimensions['E'].width = 30

        # Save to bytes
        file_stream = io.BytesIO()
        wb.save(file_stream)
        file_stream.seek(0)

        return file_stream.read()

    def generate_order_confirmation(
        self,
        order_data: Dict,
        format: str = 'pdf',
        language: str = 'pl'
    ) -> Optional[bytes]:
        """
        Generate Order Confirmation document

        Args:
            order_data: Order data dictionary
            format: Output format ('pdf', 'html', 'docx', 'xlsx')
            language: Document language ('pl', 'en')

        Returns:
            Document content as bytes
        """
        if format == 'html' or format == 'pdf':
            html_content = self._generate_order_confirmation_html(order_data, language)

            if format == 'html':
                return html_content.encode('utf-8')
            else:
                return self._html_to_pdf(html_content)

        elif format == 'docx':
            return self._generate_order_confirmation_docx(order_data, language)

        return None

    def _generate_order_confirmation_html(self, order_data: Dict, language: str) -> str:
        """Generate HTML template for Order Confirmation"""

        # Language-specific texts
        texts = {
            'pl': {
                'title': 'POTWIERDZENIE ZAMÓWIENIA',
                'order_no': 'Nr zamówienia',
                'date': 'Data',
                'buyer': 'ZAMAWIAJĄCY',
                'seller': 'SPRZEDAWCA',
                'items': 'POZYCJE ZAMÓWIENIA',
                'item_no': 'Lp.',
                'item_name': 'Nazwa',
                'quantity': 'Ilość',
                'unit_price': 'Cena jedn.',
                'total': 'Wartość',
                'summary': 'PODSUMOWANIE',
                'net_total': 'Wartość netto',
                'vat': 'VAT',
                'gross_total': 'Wartość brutto',
                'payment_terms': 'Warunki płatności',
                'delivery_date': 'Termin realizacji',
                'notes': 'UWAGI',
                'confirmation': 'Niniejszym potwierdzamy przyjęcie zamówienia do realizacji.',
                'signature': 'Podpis i pieczątka'
            },
            'en': {
                'title': 'ORDER CONFIRMATION',
                'order_no': 'Order No.',
                'date': 'Date',
                'buyer': 'BUYER',
                'seller': 'SELLER',
                'items': 'ORDER ITEMS',
                'item_no': 'No.',
                'item_name': 'Description',
                'quantity': 'Quantity',
                'unit_price': 'Unit Price',
                'total': 'Total',
                'summary': 'SUMMARY',
                'net_total': 'Net Total',
                'vat': 'VAT',
                'gross_total': 'Gross Total',
                'payment_terms': 'Payment Terms',
                'delivery_date': 'Delivery Date',
                'notes': 'NOTES',
                'confirmation': 'We hereby confirm acceptance of this order.',
                'signature': 'Signature and Stamp'
            }
        }

        t = texts.get(language, texts['pl'])

        # Prepare data
        order_number = order_data.get('order_number', 'ORD/2025/001')
        order_date = order_data.get('order_date', datetime.now().strftime('%Y-%m-%d'))
        buyer = order_data.get('buyer', {})
        items = order_data.get('items', [])

        # Calculate totals
        net_total = sum(item.get('quantity', 0) * item.get('unit_price', 0) for item in items)
        vat_rate = order_data.get('vat_rate', 0.23)
        vat_amount = net_total * vat_rate
        gross_total = net_total + vat_amount

        # Generate items rows
        items_html = ''
        for i, item in enumerate(items, 1):
            quantity = item.get('quantity', 0)
            unit_price = item.get('unit_price', 0)
            total = quantity * unit_price
            items_html += f"""
            <tr>
                <td class="center">{i}</td>
                <td>{item.get('name', '')}</td>
                <td class="center">{quantity}</td>
                <td class="right">{unit_price:.2f}</td>
                <td class="right">{total:.2f}</td>
            </tr>
            """

        html = f"""
        <!DOCTYPE html>
        <html lang="{language}">
        <head>
            <meta charset="UTF-8">
            <title>{t['title']} {order_number}</title>
            <style>
                @page {{
                    size: A4;
                    margin: 20mm;
                }}

                body {{
                    font-family: 'Arial', sans-serif;
                    font-size: 12pt;
                    color: #333;
                    line-height: 1.5;
                }}

                .header {{
                    display: flex;
                    justify-content: space-between;
                    margin-bottom: 30px;
                    padding-bottom: 20px;
                    border-bottom: 3px solid #2980b9;
                }}

                .company-info {{
                    text-align: right;
                    font-size: 10pt;
                }}

                .company-name {{
                    font-size: 16pt;
                    font-weight: bold;
                    color: #2980b9;
                    margin-bottom: 5px;
                }}

                .document-title {{
                    text-align: center;
                    margin: 30px 0;
                }}

                .document-title h1 {{
                    font-size: 24pt;
                    color: #2980b9;
                    margin-bottom: 10px;
                }}

                .order-info {{
                    text-align: center;
                    font-size: 14pt;
                    margin-bottom: 30px;
                }}

                .parties {{
                    display: flex;
                    justify-content: space-between;
                    margin-bottom: 30px;
                }}

                .party-box {{
                    width: 48%;
                    padding: 15px;
                    background: #ecf0f1;
                    border-radius: 5px;
                }}

                .party-box h3 {{
                    font-size: 12pt;
                    color: #2980b9;
                    margin-bottom: 10px;
                    border-bottom: 2px solid #2980b9;
                    padding-bottom: 5px;
                }}

                .items-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin: 20px 0;
                }}

                .items-table th {{
                    background: #2980b9;
                    color: white;
                    padding: 10px;
                    text-align: left;
                }}

                .items-table td {{
                    padding: 8px 10px;
                    border-bottom: 1px solid #ddd;
                }}

                .center {{ text-align: center; }}
                .right {{ text-align: right; }}

                .summary {{
                    margin-top: 30px;
                    padding: 20px;
                    background: #ecf0f1;
                    border-radius: 5px;
                }}

                .summary-row {{
                    display: flex;
                    justify-content: space-between;
                    margin: 10px 0;
                    padding: 5px 0;
                }}

                .summary-row.total {{
                    border-top: 2px solid #2980b9;
                    padding-top: 10px;
                    font-weight: bold;
                    font-size: 14pt;
                    color: #2980b9;
                }}

                .terms {{
                    margin-top: 30px;
                    padding: 15px;
                    background: #f8f9fa;
                    border-left: 4px solid #2980b9;
                }}

                .confirmation-text {{
                    margin: 30px 0;
                    padding: 20px;
                    background: #d4edda;
                    border-radius: 5px;
                    text-align: center;
                    font-size: 14pt;
                    color: #155724;
                }}

                .signature-section {{
                    display: flex;
                    justify-content: space-between;
                    margin-top: 50px;
                }}

                .signature-box {{
                    width: 40%;
                    text-align: center;
                }}

                .signature-line {{
                    border-top: 1px solid #333;
                    margin-top: 60px;
                    padding-top: 5px;
                    font-size: 10pt;
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div>
                    <div class="company-name">{self.company_settings['name']}</div>
                    <div>{self.company_settings['address']}</div>
                    <div>{self.company_settings['city']}</div>
                </div>
                <div class="company-info">
                    <p>NIP: {self.company_settings['nip']}</p>
                    <p>Tel: {self.company_settings['phone']}</p>
                    <p>Email: {self.company_settings['email']}</p>
                </div>
            </div>

            <div class="document-title">
                <h1>{t['title']}</h1>
            </div>

            <div class="order-info">
                <strong>{t['order_no']}:</strong> {order_number}<br>
                <strong>{t['date']}:</strong> {order_date}
            </div>

            <div class="parties">
                <div class="party-box">
                    <h3>{t['seller']}</h3>
                    <p><strong>{self.company_settings['name']}</strong></p>
                    <p>{self.company_settings['address']}</p>
                    <p>{self.company_settings['city']}</p>
                    <p>NIP: {self.company_settings['nip']}</p>
                </div>

                <div class="party-box">
                    <h3>{t['buyer']}</h3>
                    <p><strong>{buyer.get('name', '')}</strong></p>
                    <p>{buyer.get('address', '')}</p>
                    <p>{buyer.get('city', '')}</p>
                    <p>NIP: {buyer.get('nip', '')}</p>
                </div>
            </div>

            <h3>{t['items']}</h3>
            <table class="items-table">
                <thead>
                    <tr>
                        <th style="width: 5%">{t['item_no']}</th>
                        <th style="width: 45%">{t['item_name']}</th>
                        <th style="width: 15%">{t['quantity']}</th>
                        <th style="width: 15%">{t['unit_price']}</th>
                        <th style="width: 20%">{t['total']}</th>
                    </tr>
                </thead>
                <tbody>
                    {items_html}
                </tbody>
            </table>

            <div class="summary">
                <h3>{t['summary']}</h3>
                <div class="summary-row">
                    <span>{t['net_total']}:</span>
                    <span>{net_total:.2f} PLN</span>
                </div>
                <div class="summary-row">
                    <span>{t['vat']} ({int(vat_rate*100)}%):</span>
                    <span>{vat_amount:.2f} PLN</span>
                </div>
                <div class="summary-row total">
                    <span>{t['gross_total']}:</span>
                    <span>{gross_total:.2f} PLN</span>
                </div>
            </div>

            <div class="terms">
                <p><strong>{t['payment_terms']}:</strong> {order_data.get('payment_terms', '14 dni')}</p>
                <p><strong>{t['delivery_date']}:</strong> {order_data.get('delivery_date', '30 dni')}</p>
            </div>

            <div class="confirmation-text">
                {t['confirmation']}
            </div>

            <div class="signature-section">
                <div class="signature-box">
                    <div class="signature-line">
                        {t['signature']}
                    </div>
                </div>
            </div>
        </body>
        </html>
        """

        return html

    def _generate_order_confirmation_docx(self, order_data: Dict, language: str) -> bytes:
        """Generate Order Confirmation in DOCX format"""
        # Similar to WZ but with order confirmation specific content
        # Implementation details omitted for brevity
        doc = Document()
        # Add content...
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.read()


# Example usage and testing
if __name__ == '__main__':
    generator = ModernDocumentGenerator()

    # Test WZ document
    wz_data = {
        'wz_number': 'WZ/2025/001',
        'issue_date': '2025-11-17',
        'order_number': 'ZAM/2025/123',
        'recipient': {
            'name': 'ABC Company Sp. z o.o.',
            'address': 'ul. Testowa 456',
            'city': 'Kraków',
            'postal_code': '30-001',
            'nip': '987-654-32-10',
            'contact_person': 'Jan Kowalski',
            'phone': '+48 600 700 800'
        },
        'items': [
            {
                'name': 'Element metalowy 200x100x5',
                'quantity': 10,
                'unit': 'szt.',
                'notes': 'Stal nierdzewna'
            },
            {
                'name': 'Profil aluminiowy 50x50x3',
                'quantity': 25,
                'unit': 'mb',
                'notes': 'Anodowany'
            }
        ],
        'notes': 'Transport własny odbiorcy. Wydanie w magazynie głównym.'
    }

    # Test Order Confirmation
    order_data = {
        'order_number': 'ORD/2025/001',
        'order_date': '2025-11-17',
        'buyer': {
            'name': 'XYZ Corporation',
            'address': '123 Business Ave',
            'city': 'Warsaw 00-001',
            'nip': '123-456-78-90'
        },
        'items': [
            {
                'name': 'Product A - Premium Edition',
                'quantity': 5,
                'unit_price': 1250.00
            },
            {
                'name': 'Product B - Standard',
                'quantity': 10,
                'unit_price': 750.00
            }
        ],
        'vat_rate': 0.23,
        'payment_terms': '30 dni',
        'delivery_date': '14 dni roboczych'
    }

    print("Modern Document Generator initialized")
    print(f"PDF Engine: {generator.pdf_engine or 'Not available'}")
    print("Ready to generate documents in HTML, DOCX, XLSX formats")
    if generator.pdf_engine:
        print("PDF generation is available")
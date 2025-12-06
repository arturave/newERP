#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WZ Generator Module
Generowanie dokumentów wydania zewnętrznego (WZ) w formatach PDF, Word i Excel
"""

import os
from datetime import datetime, date
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, asdict
import json

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter


@dataclass
class WZItem:
    """Pozycja dokumentu WZ"""
    lp: int  # Lp.
    name: str  # Nazwa produktu/części
    quantity: int  # Ilość
    unit: str = "szt"  # Jednostka miary
    notes: str = ""  # Uwagi

    def to_dict(self) -> Dict:
        return asdict(self)

    @classmethod
    def from_dict(cls, data: Dict) -> 'WZItem':
        return cls(**data)


@dataclass
class WZData:
    """Dane dokumentu WZ"""
    wz_number: str
    order_id: str
    process_no: str
    issue_date: date

    # Dane odbiorcy
    recipient_name: str
    recipient_address: str = ""
    recipient_city: str = ""
    recipient_postal_code: str = ""
    recipient_nip: str = ""
    recipient_contact_person: str = ""
    recipient_contact_phone: str = ""

    # Pozycje
    items: List[WZItem] = None

    # Dodatkowe informacje
    notes: str = ""
    transport_info: str = ""

    # Dane wystawcy (firma)
    issuer_name: str = "Twoja Firma Sp. z o.o."
    issuer_address: str = "ul. Produkcyjna 123"
    issuer_city: str = "00-000 Miasto"
    issuer_nip: str = "123-456-78-90"
    issuer_phone: str = "+48 123 456 789"
    issuer_email: str = "biuro@firma.pl"

    def __post_init__(self):
        if self.items is None:
            self.items = []


class WZGenerator:
    """Generator dokumentów WZ w różnych formatach"""

    def __init__(self, db_client):
        """
        Inicjalizacja generatora WZ

        Args:
            db_client: Klient Supabase
        """
        self.client = db_client

    def get_wz_number(self, order_id: str) -> str:
        """
        Generuje numer WZ na podstawie numeru zamówienia

        Args:
            order_id: ID zamówienia

        Returns:
            Numer WZ w formacie WZ-{process_no}
        """
        try:
            # Pobierz numer procesowy zamówienia
            response = self.client.table('orders').select(
                'process_no'
            ).eq('id', order_id).execute()

            if not response.data:
                raise ValueError(f"Zamówienie {order_id} nie zostało znalezione")

            process_no = response.data[0]['process_no']
            return f"WZ-{process_no}"

        except Exception as e:
            print(f"❌ Błąd generowania numeru WZ: {e}")
            return None

    def get_order_data_for_wz(self, order_id: str) -> Optional[WZData]:
        """
        Pobiera dane zamówienia i tworzy obiekt WZData

        Args:
            order_id: ID zamówienia

        Returns:
            WZData lub None w przypadku błędu
        """
        try:
            # Pobierz dane zamówienia i klienta
            response = self.client.table('v_orders_full').select('*').eq(
                'id', order_id
            ).execute()

            if not response.data:
                print(f"❌ Zamówienie {order_id} nie zostało znalezione")
                return None

            order = response.data[0]

            # Pobierz części zamówienia
            parts_response = self.client.table('parts').select('*').eq(
                'order_id', order_id
            ).order('name').execute()

            # Utwórz pozycje WZ z części
            items = []
            for i, part in enumerate(parts_response.data, 1):
                item = WZItem(
                    lp=i,
                    name=part['name'],
                    quantity=part.get('qty', 1),
                    unit='szt',
                    notes=f"{part.get('material', '')} {part.get('thickness_mm', '')}mm".strip()
                )
                items.append(item)

            # Utwórz obiekt WZData
            wz_data = WZData(
                wz_number=self.get_wz_number(order_id),
                order_id=order_id,
                process_no=order['process_no'],
                issue_date=date.today(),
                recipient_name=order.get('customer_name', ''),
                recipient_address=order.get('customer_address', ''),
                recipient_city=order.get('customer_city', ''),
                recipient_postal_code=order.get('customer_postal_code', ''),
                recipient_nip=order.get('customer_nip', ''),
                recipient_contact_person=order.get('customer_contact_person', ''),
                recipient_contact_phone=order.get('customer_contact_phone', ''),
                items=items,
                notes=order.get('notes', '')
            )

            return wz_data

        except Exception as e:
            print(f"❌ Błąd pobierania danych zamówienia: {e}")
            return None

    def generate_pdf(
        self,
        wz_data: WZData,
        output_path: str,
        logo_path: Optional[str] = None
    ) -> bool:
        """
        Generuje dokument WZ w formacie PDF

        Args:
            wz_data: Dane dokumentu WZ
            output_path: Ścieżka do pliku wyjściowego
            logo_path: Opcjonalna ścieżka do logo firmy

        Returns:
            True jeśli sukces, False w przypadku błędu
        """
        try:
            # Utwórz dokument PDF
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                topMargin=1.5*cm,
                bottomMargin=1.5*cm,
                leftMargin=2*cm,
                rightMargin=2*cm
            )

            # Style
            styles = getSampleStyleSheet()

            # Styl nagłówka
            style_header = ParagraphStyle(
                'CustomHeader',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#1F4E79'),
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName='Helvetica-Bold'
            )

            # Styl zwykłego tekstu
            style_normal = styles['Normal']
            style_normal.fontSize = 10

            # Styl małego tekstu
            style_small = ParagraphStyle(
                'Small',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#666666')
            )

            # Elementy dokumentu
            elements = []

            # Logo (jeśli podane)
            if logo_path and os.path.exists(logo_path):
                logo = RLImage(logo_path, width=3*cm, height=2*cm)
                elements.append(logo)
                elements.append(Spacer(1, 0.5*cm))

            # Nagłówek
            elements.append(Paragraph("DOKUMENT WYDANIA ZEWNĘTRZNEGO", style_header))
            elements.append(Spacer(1, 0.3*cm))

            # Informacje o dokumencie
            doc_info = f"""
            <para align=center fontSize=12>
            <b>WZ Nr: {wz_data.wz_number}</b><br/>
            Data wystawienia: {wz_data.issue_date.strftime('%Y-%m-%d')}<br/>
            Zamówienie: {wz_data.process_no}
            </para>
            """
            elements.append(Paragraph(doc_info, style_normal))
            elements.append(Spacer(1, 0.8*cm))

            # Tabela: Wystawca | Odbiorca
            header_data = [
                ['WYSTAWCA', 'ODBIORCA']
            ]

            issuer_text = f"""
            <b>{wz_data.issuer_name}</b><br/>
            {wz_data.issuer_address}<br/>
            {wz_data.issuer_city}<br/>
            NIP: {wz_data.issuer_nip}<br/>
            Tel: {wz_data.issuer_phone}<br/>
            Email: {wz_data.issuer_email}
            """

            recipient_text = f"""
            <b>{wz_data.recipient_name}</b><br/>
            {wz_data.recipient_address}<br/>
            {wz_data.recipient_postal_code} {wz_data.recipient_city}<br/>
            {f'NIP: {wz_data.recipient_nip}<br/>' if wz_data.recipient_nip else ''}
            {f'Osoba kontaktowa: {wz_data.recipient_contact_person}<br/>' if wz_data.recipient_contact_person else ''}
            {f'Tel: {wz_data.recipient_contact_phone}' if wz_data.recipient_contact_phone else ''}
            """

            body_data = [
                [Paragraph(issuer_text, style_normal), Paragraph(recipient_text, style_normal)]
            ]

            header_table = Table(header_data + body_data, colWidths=[8*cm, 8*cm])
            header_table.setStyle(TableStyle([
                # Nagłówek
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('TOPPADDING', (0, 0), (-1, 0), 10),

                # Ciało
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#F5F5F5')),
                ('VALIGN', (0, 1), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 1), (-1, -1), 10),
                ('RIGHTPADDING', (0, 1), (-1, -1), 10),
                ('TOPPADDING', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 10),

                # Obramowanie
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            elements.append(header_table)
            elements.append(Spacer(1, 0.8*cm))

            # Tabela pozycji
            items_header = ['Lp.', 'Nazwa', 'Ilość', 'Jednostka', 'Uwagi']
            items_data = [items_header]

            for item in wz_data.items:
                items_data.append([
                    str(item.lp),
                    item.name,
                    str(item.quantity),
                    item.unit,
                    item.notes
                ])

            items_table = Table(items_data, colWidths=[1.5*cm, 7*cm, 2*cm, 2*cm, 3.5*cm])
            items_table.setStyle(TableStyle([
                # Nagłówek
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('VALIGN', (0, 0), (-1, 0), 'MIDDLE'),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),

                # Ciało
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('ALIGN', (0, 1), (0, -1), 'CENTER'),  # Lp.
                ('ALIGN', (2, 1), (3, -1), 'CENTER'),  # Ilość, Jednostka
                ('VALIGN', (0, 1), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 1), (-1, -1), 6),
                ('RIGHTPADDING', (0, 1), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),

                # Obramowanie
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
            ]))

            elements.append(items_table)
            elements.append(Spacer(1, 0.5*cm))

            # Podsumowanie
            summary_text = f"<b>Razem pozycji: {len(wz_data.items)}</b>"
            elements.append(Paragraph(summary_text, style_normal))
            elements.append(Spacer(1, 0.5*cm))

            # Uwagi
            if wz_data.notes:
                elements.append(Paragraph(f"<b>Uwagi:</b>", style_normal))
                elements.append(Paragraph(wz_data.notes, style_normal))
                elements.append(Spacer(1, 0.3*cm))

            # Transport
            if wz_data.transport_info:
                elements.append(Paragraph(f"<b>Transport:</b>", style_normal))
                elements.append(Paragraph(wz_data.transport_info, style_normal))
                elements.append(Spacer(1, 0.3*cm))

            elements.append(Spacer(1, 1*cm))

            # Podpisy
            signatures = Table([
                ['Wydał:', 'Odebrał:'],
                ['', ''],
                ['................................', '................................'],
                ['(podpis)', '(podpis)']
            ], colWidths=[8*cm, 8*cm])

            signatures.setStyle(TableStyle([
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
                ('ALIGN', (0, 2), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'BOTTOM'),
                ('TOPPADDING', (0, 2), (-1, 2), 20),
                ('BOTTOMPADDING', (0, 2), (-1, 2), 0),
            ]))

            elements.append(signatures)
            elements.append(Spacer(1, 0.5*cm))

            # Stopka
            footer_text = f"""
            <para align=center fontSize=8 textColor=#666666>
            Dokument wygenerowany automatycznie dnia {datetime.now().strftime('%Y-%m-%d %H:%M')}<br/>
            System Zarządzania Produkcją
            </para>
            """
            elements.append(Paragraph(footer_text, style_small))

            # Zbuduj dokument
            doc.build(elements)

            print(f"✅ Wygenerowano PDF: {output_path}")
            return True

        except Exception as e:
            print(f"❌ Błąd generowania PDF: {e}")
            import traceback
            traceback.print_exc()
            return False

    def generate_word(self, wz_data: WZData, output_path: str) -> bool:
        """
        Generuje dokument WZ w formacie Word (DOCX)

        Args:
            wz_data: Dane dokumentu WZ
            output_path: Ścieżka do pliku wyjściowego

        Returns:
            True jeśli sukces, False w przypadku błędu
        """
        try:
            # Utwórz dokument Word
            doc = Document()

            # Nagłówek
            heading = doc.add_heading('DOKUMENT WYDANIA ZEWNĘTRZNEGO', 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Informacje o dokumencie
            doc_info = doc.add_paragraph()
            doc_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc_info.add_run(f"WZ Nr: {wz_data.wz_number}\n").bold = True
            doc_info.add_run(f"Data wystawienia: {wz_data.issue_date.strftime('%Y-%m-%d')}\n")
            doc_info.add_run(f"Zamówienie: {wz_data.process_no}")

            doc.add_paragraph()  # Odstęp

            # Tabela: Wystawca | Odbiorca
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Light Grid Accent 1'

            # Nagłówki
            header_cells = table.rows[0].cells
            header_cells[0].text = 'WYSTAWCA'
            header_cells[1].text = 'ODBIORCA'

            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Dane
            issuer_cell = table.rows[1].cells[0]
            issuer_text = f"{wz_data.issuer_name}\n{wz_data.issuer_address}\n{wz_data.issuer_city}\n"
            issuer_text += f"NIP: {wz_data.issuer_nip}\nTel: {wz_data.issuer_phone}\nEmail: {wz_data.issuer_email}"
            issuer_cell.text = issuer_text

            recipient_cell = table.rows[1].cells[1]
            recipient_text = f"{wz_data.recipient_name}\n{wz_data.recipient_address}\n"
            recipient_text += f"{wz_data.recipient_postal_code} {wz_data.recipient_city}\n"
            if wz_data.recipient_nip:
                recipient_text += f"NIP: {wz_data.recipient_nip}\n"
            if wz_data.recipient_contact_person:
                recipient_text += f"Kontakt: {wz_data.recipient_contact_person}\n"
            if wz_data.recipient_contact_phone:
                recipient_text += f"Tel: {wz_data.recipient_contact_phone}"
            recipient_cell.text = recipient_text

            doc.add_paragraph()  # Odstęp

            # Tabela pozycji
            doc.add_heading('Pozycje:', level=2)

            items_table = doc.add_table(rows=1 + len(wz_data.items), cols=5)
            items_table.style = 'Light Grid Accent 1'

            # Nagłówki
            header_cells = items_table.rows[0].cells
            headers = ['Lp.', 'Nazwa', 'Ilość', 'Jednostka', 'Uwagi']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Pozycje
            for i, item in enumerate(wz_data.items, 1):
                row = items_table.rows[i]
                row.cells[0].text = str(item.lp)
                row.cells[1].text = item.name
                row.cells[2].text = str(item.quantity)
                row.cells[3].text = item.unit
                row.cells[4].text = item.notes

            # Podsumowanie
            doc.add_paragraph()
            summary = doc.add_paragraph()
            summary.add_run(f"Razem pozycji: {len(wz_data.items)}").bold = True

            # Uwagi
            if wz_data.notes:
                doc.add_paragraph()
                doc.add_paragraph("Uwagi:").runs[0].bold = True
                doc.add_paragraph(wz_data.notes)

            # Transport
            if wz_data.transport_info:
                doc.add_paragraph()
                doc.add_paragraph("Transport:").runs[0].bold = True
                doc.add_paragraph(wz_data.transport_info)

            # Podpisy
            doc.add_page_break()
            doc.add_paragraph()
            doc.add_paragraph()

            sig_table = doc.add_table(rows=2, cols=2)
            sig_table.rows[0].cells[0].text = "Wydał:"
            sig_table.rows[0].cells[1].text = "Odebrał:"
            sig_table.rows[1].cells[0].text = "\n\n................................\n(podpis)"
            sig_table.rows[1].cells[1].text = "\n\n................................\n(podpis)"

            for row in sig_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Zapisz dokument
            doc.save(output_path)

            print(f"✅ Wygenerowano Word: {output_path}")
            return True

        except Exception as e:
            print(f"❌ Błąd generowania Word: {e}")
            import traceback
            traceback.print_exc()
            return False

    def generate_excel(self, wz_data: WZData, output_path: str) -> bool:
        """
        Generuje dokument WZ w formacie Excel (XLSX)

        Args:
            wz_data: Dane dokumentu WZ
            output_path: Ścieżka do pliku wyjściowego

        Returns:
            True jeśli sukces, False w przypadku błędu
        """
        try:
            # Utwórz workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "WZ"

            # Style
            header_font = Font(bold=True, size=14, color="FFFFFF")
            header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
            normal_font = Font(size=11)
            bold_font = Font(bold=True, size=11)

            border_thin = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Szerokości kolumn
            ws.column_dimensions['A'].width = 5
            ws.column_dimensions['B'].width = 40
            ws.column_dimensions['C'].width = 10
            ws.column_dimensions['D'].width = 12
            ws.column_dimensions['E'].width = 30

            row = 1

            # Nagłówek dokumentu
            ws.merge_cells(f'A{row}:E{row}')
            cell = ws[f'A{row}']
            cell.value = "DOKUMENT WYDANIA ZEWNĘTRZNEGO"
            cell.font = Font(bold=True, size=16)
            cell.alignment = Alignment(horizontal='center', vertical='center')
            row += 2

            # Informacje o dokumencie
            ws[f'A{row}'] = "Nr WZ:"
            ws[f'A{row}'].font = bold_font
            ws[f'B{row}'] = wz_data.wz_number
            row += 1

            ws[f'A{row}'] = "Data:"
            ws[f'A{row}'].font = bold_font
            ws[f'B{row}'] = wz_data.issue_date.strftime('%Y-%m-%d')
            row += 1

            ws[f'A{row}'] = "Zamówienie:"
            ws[f'A{row}'].font = bold_font
            ws[f'B{row}'] = wz_data.process_no
            row += 2

            # Wystawca
            ws[f'A{row}'] = "WYSTAWCA"
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].fill = header_fill
            row += 1
            ws[f'A{row}'] = wz_data.issuer_name
            ws[f'A{row}'].font = bold_font
            row += 1
            ws[f'A{row}'] = wz_data.issuer_address
            row += 1
            ws[f'A{row}'] = wz_data.issuer_city
            row += 1
            ws[f'A{row}'] = f"NIP: {wz_data.issuer_nip}"
            row += 2

            # Odbiorca
            ws[f'A{row}'] = "ODBIORCA"
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].fill = header_fill
            row += 1
            ws[f'A{row}'] = wz_data.recipient_name
            ws[f'A{row}'].font = bold_font
            row += 1
            ws[f'A{row}'] = wz_data.recipient_address
            row += 1
            ws[f'A{row}'] = f"{wz_data.recipient_postal_code} {wz_data.recipient_city}"
            row += 1
            if wz_data.recipient_nip:
                ws[f'A{row}'] = f"NIP: {wz_data.recipient_nip}"
                row += 1
            row += 1

            # Tabela pozycji - nagłówki
            headers = ['Lp.', 'Nazwa', 'Ilość', 'Jednostka', 'Uwagi']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=row, column=col)
                cell.value = header
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = border_thin

            row += 1

            # Pozycje
            for item in wz_data.items:
                ws.cell(row=row, column=1, value=item.lp).border = border_thin
                ws.cell(row=row, column=2, value=item.name).border = border_thin
                ws.cell(row=row, column=3, value=item.quantity).border = border_thin
                ws.cell(row=row, column=4, value=item.unit).border = border_thin
                ws.cell(row=row, column=5, value=item.notes).border = border_thin

                # Wyśrodkowanie liczb
                ws.cell(row=row, column=1).alignment = Alignment(horizontal='center')
                ws.cell(row=row, column=3).alignment = Alignment(horizontal='center')
                ws.cell(row=row, column=4).alignment = Alignment(horizontal='center')

                row += 1

            row += 1

            # Podsumowanie
            ws[f'A{row}'] = f"Razem pozycji:"
            ws[f'A{row}'].font = bold_font
            ws[f'B{row}'] = len(wz_data.items)
            ws[f'B{row}'].font = bold_font
            row += 2

            # Uwagi
            if wz_data.notes:
                ws[f'A{row}'] = "Uwagi:"
                ws[f'A{row}'].font = bold_font
                row += 1
                ws[f'A{row}'] = wz_data.notes
                row += 2

            # Transport
            if wz_data.transport_info:
                ws[f'A{row}'] = "Transport:"
                ws[f'A{row}'].font = bold_font
                row += 1
                ws[f'A{row}'] = wz_data.transport_info
                row += 2

            row += 2

            # Podpisy
            ws[f'A{row}'] = "Wydał:"
            ws[f'C{row}'] = "Odebrał:"
            row += 3
            ws[f'A{row}'] = "................................"
            ws[f'C{row}'] = "................................"
            row += 1
            ws[f'A{row}'] = "(podpis)"
            ws[f'C{row}'] = "(podpis)"

            # Zapisz workbook
            wb.save(output_path)

            print(f"✅ Wygenerowano Excel: {output_path}")
            return True

        except Exception as e:
            print(f"❌ Błąd generowania Excel: {e}")
            import traceback
            traceback.print_exc()
            return False

    def save_wz_to_db(self, wz_data: WZData) -> Optional[Dict]:
        """
        Zapisuje dokument WZ do bazy danych

        Args:
            wz_data: Dane dokumentu WZ

        Returns:
            Dict z danymi zapisanego WZ lub None w przypadku błędu
        """
        try:
            # Konwertuj items do JSON
            items_json = [item.to_dict() for item in wz_data.items]

            # Dane do zapisu
            db_data = {
                'wz_number': wz_data.wz_number,
                'order_id': wz_data.order_id,
                'issue_date': wz_data.issue_date.isoformat(),
                'recipient_name': wz_data.recipient_name,
                'recipient_address': wz_data.recipient_address,
                'recipient_city': wz_data.recipient_city,
                'recipient_postal_code': wz_data.recipient_postal_code,
                'recipient_nip': wz_data.recipient_nip,
                'recipient_contact_person': wz_data.recipient_contact_person,
                'recipient_contact_phone': wz_data.recipient_contact_phone,
                'items': json.dumps(items_json),
                'notes': wz_data.notes,
                'transport_info': wz_data.transport_info,
                'status': 'ISSUED'
            }

            response = self.client.table('delivery_notes').insert(db_data).execute()

            if response.data:
                print(f"✅ Zapisano WZ do bazy: {wz_data.wz_number}")
                return response.data[0]

            return None

        except Exception as e:
            print(f"❌ Błąd zapisu WZ do bazy: {e}")
            return None


# Przykład użycia
if __name__ == '__main__':
    print("WZGenerator - Generator dokumentów WZ")
    print("=" * 50)
    print("Obsługiwane formaty:")
    print("- PDF (profesjonalny, z tabelami)")
    print("- Word (DOCX, edytowalny)")
    print("- Excel (XLSX, arkusz kalkulacyjny)")
